"""
Flask-based web interface for LitOrganizer.

Provides a local web server with real-time progress updates via WebSocket.
"""

import os
import re
import gc
import json
import time
import logging
import threading
import traceback
import webbrowser
from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from flask_socketio import SocketIO, emit

from modules.core.pdf_renamer import PDFProcessor
from modules.utils.file_utils import ensure_dir, get_version
from modules.utils.pdf_metadata_extractor import load_api_config, extract_doi

# ---------------------------------------------------------------------------
# Application factory
# ---------------------------------------------------------------------------

def create_app() -> tuple:
    """Create and configure the Flask application with SocketIO."""
    app_root = Path(__file__).resolve().parent.parent.parent
    
    app = Flask(
        __name__,
        template_folder=str(Path(__file__).parent / 'templates'),
        static_folder=str(Path(__file__).parent / 'static'),
    )
    app.config['SECRET_KEY'] = 'litorganizer-local-key'
    app.config['APP_ROOT'] = str(app_root)
    app.config['TEMPLATES_AUTO_RELOAD'] = True
    app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
    app.jinja_env.auto_reload = True
    
    socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')
    
    # Serve the resources folder (logos, icons, etc.)
    @app.route('/resources/<path:filename>')
    def serve_resource(filename):
        return send_from_directory(str(app_root / 'resources'), filename)
    
    # -----------------------------------------------------------------------
    # Shared state
    # -----------------------------------------------------------------------
    state: Dict[str, Any] = {
        'processing': False,
        'searching': False,
        'worker_thread': None,
        'search_thread': None,
        'stop_flag': False,
        'search_stop_flag': False,
        'log_messages': [],
        'file_statuses': [],
        'search_results': [],
        'last_stats': None,
        'last_completed_stats': None,
        'process_start_time': None,
        'process_end_time': None,
        'search_progress': 0,
        'process_progress': 0,
        'process_directory': '',
        'process_total_files': 0,
    }
    
    # -----------------------------------------------------------------------
    # Custom SocketIO log handler
    # -----------------------------------------------------------------------
    class SocketIOLogHandler(logging.Handler):
        """Forwards log records to connected WebSocket clients."""
        def emit(self, record):
            msg = self.format(record)
            state['log_messages'].append(msg)
            try:
                socketio.emit('log_message', {'message': msg})
            except Exception:
                pass
    
    # -----------------------------------------------------------------------
    # HTTP Routes
    # -----------------------------------------------------------------------
    
    @app.route('/')
    def index():
        """Main page – PDF processing."""
        return render_template('index.html', version=get_version())
    
    @app.route('/guide')
    def guide_page():
        """Usage guide page."""
        return render_template('guide.html', version=get_version())
    
    @app.route('/search')
    def search_page():
        """Keyword search page."""
        # Pass existing search results to the template
        existing_results = state.get('search_results', [])
        return render_template('search.html', version=get_version(), 
                               existing_results=existing_results,
                               result_count=len(existing_results))
    
    @app.route('/api/get_search_results')
    def get_search_results():
        """Return current search results."""
        results = state.get('search_results', [])
        return jsonify({
            'results': results,
            'count': len(results),
            'searching': state.get('searching', False)
        })
    
    @app.route('/api/status')
    def get_global_status():
        """Return global status of all active processes."""
        import time
        return jsonify({
            'processing': state.get('processing', False),
            'searching': state.get('searching', False),
            'search_results_count': len(state.get('search_results', [])),
            'process_start_time': state.get('process_start_time'),
            'search_progress': state.get('search_progress', 0),
            'process_progress': state.get('process_progress', 0),
        })
    
    @app.route('/api/process_status')
    def get_process_status():
        """Return current processing state including file statuses for page reload persistence."""
        return jsonify({
            'processing': state.get('processing', False),
            'file_statuses': state.get('file_statuses', []),
            'process_progress': state.get('process_progress', 0),
            'process_directory': state.get('process_directory', ''),
            'total_files': state.get('process_total_files', 0),
            'last_completed_stats': state.get('last_completed_stats'),
        })
    
    @app.route('/api/open_folder', methods=['POST'])
    def open_folder():
        """Open a folder in the system file explorer."""
        data = request.get_json() or {}
        folder_path = data.get('path', '')
        
        if not folder_path:
            return jsonify({'success': False, 'message': 'No path provided.'}), 400
        
        p = Path(folder_path)
        if not p.exists() or not p.is_dir():
            return jsonify({'success': False, 'message': 'Directory not found.'}), 404
        
        try:
            import subprocess
            import platform
            system = platform.system()
            if system == 'Windows':
                os.startfile(str(p))
            elif system == 'Darwin':
                subprocess.Popen(['open', str(p)])
            else:
                subprocess.Popen(['xdg-open', str(p)])
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500
    
    @app.route('/api/native_browse', methods=['POST'])
    def native_browse():
        """Open native OS folder picker dialog via tkinter."""
        import platform as _platform
        result = {'path': ''}

        def pick_folder():
            try:
                import tkinter as tk
                from tkinter import filedialog
                root = tk.Tk()
                root.withdraw()
                if _platform.system() == 'Windows':
                    root.wm_attributes('-topmost', 1)
                folder = filedialog.askdirectory(title="Select PDF Directory")
                root.destroy()
                result['path'] = folder or ''
            except Exception as e:
                logging.error(f"Native browse error: {e}")
                result['path'] = ''

        t = threading.Thread(target=pick_folder)
        t.start()
        t.join(timeout=120)

        if result['path']:
            return jsonify({'success': True, 'path': result['path']})
        return jsonify({'success': False, 'message': 'No folder selected'})

    @app.route('/api/quick_paths')
    def quick_paths():
        """Return common quick-access directory paths for the file browser modal."""
        import platform as _platform
        home = Path.home()
        paths = [
            {'name': 'Desktop', 'icon': 'desktop', 'path': str(home / 'Desktop')},
            {'name': 'Documents', 'icon': 'folder', 'path': str(home / 'Documents')},
            {'name': 'Downloads', 'icon': 'download', 'path': str(home / 'Downloads')},
        ]
        if _platform.system() == 'Windows':
            import string
            for letter in string.ascii_uppercase:
                drive = f"{letter}:\\"
                if Path(drive).exists():
                    paths.append({'name': f'Drive ({letter}:)', 'icon': 'drive', 'path': drive})
        return jsonify(paths)

    @app.route('/statistics')
    def statistics_page():
        """Statistics page."""
        stats = state.get('last_stats')
        return render_template('statistics.html', version=get_version(), stats=stats)
    
    @app.route('/settings')
    def settings_page():
        """API settings page."""
        config = load_api_config()
        return render_template('settings.html', version=get_version(), config=config)
    
    @app.route('/settings/save', methods=['POST'])
    def save_settings():
        """Save API settings to config/api_keys.json."""
        try:
            data = request.get_json()
            config_path = Path(app.config['APP_ROOT']) / 'config' / 'api_keys.json'
            config_path.parent.mkdir(parents=True, exist_ok=True)
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
            return jsonify({'success': True, 'message': 'Settings saved successfully.'})
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500
    
    @app.route('/api/browse', methods=['POST'])
    def browse_directory():
        """Return sub-directories of a given path for the directory browser."""
        data = request.get_json() or {}
        path_str = data.get('path', '')
        
        if not path_str:
            # Default to the application root directory
            path_str = app.config['APP_ROOT']
        
        p = Path(path_str)
        if not p.exists() or not p.is_dir():
            return jsonify({'dirs': [], 'current': path_str, 'error': 'Directory not found'}), 404
        
        try:
            dirs = sorted([
                str(d) for d in p.iterdir()
                if d.is_dir() and not d.name.startswith('.')
            ])
        except PermissionError:
            dirs = []
        
        parent = str(p.parent) if str(p.parent) != str(p) else ''
        return jsonify({'dirs': dirs, 'current': str(p), 'parent': parent})
    
    @app.route('/api/validate_dir', methods=['POST'])
    def validate_directory():
        """Check if a directory exists and contains PDF files."""
        data = request.get_json() or {}
        path_str = data.get('path', '')
        p = Path(path_str)
        
        if not p.exists() or not p.is_dir():
            return jsonify({'valid': False, 'message': 'Directory does not exist.'})
        
        pdf_count = len(list(p.glob('*.pdf')))
        return jsonify({
            'valid': True,
            'pdf_count': pdf_count,
            'message': f'{pdf_count} PDF files found.' if pdf_count else 'No PDF files found in this directory.'
        })
    
    @app.route('/api/download_search_results', methods=['POST'])
    def download_search_results():
        """Export search results as Excel or Word file."""
        data = request.get_json() or {}
        fmt = data.get('format', 'xlsx')
        filename = data.get('filename', 'search_results')
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        results = state.get('search_results', [])
        
        if not results:
            return jsonify({'error': 'No results to export'}), 400
        
        export_dir = Path(app.config['APP_ROOT']) / 'exports'
        export_dir.mkdir(exist_ok=True)
        
        if fmt == 'xlsx':
            import pandas as pd
            df = pd.DataFrame(results, columns=[
                'DOI', 'PDF Name', 'Page', 'Keyword',
                'Previous Sentence', 'Matched Sentence', 'Next Sentence'
            ])
            out_path = export_dir / f'{filename}.xlsx'
            df.to_excel(str(out_path), index=False)
            return send_file(str(out_path), as_attachment=True, download_name=f'{filename}.xlsx')
        
        elif fmt == 'docx':
            from docx import Document
            from docx.shared import RGBColor
            from docx.enum.text import WD_COLOR_INDEX
            import re as regex
            
            doc = Document()
            doc.add_heading(f'Search Results', 0)
            
            for row in results:
                doi, fname, page, kw, prev_s, match_s, next_s = row
                if doi:
                    p = doc.add_paragraph()
                    p.add_run(f'DOI: {doi}').bold = True
                h = doc.add_heading(level=1)
                h.add_run(f'File: {fname}')
                pp = doc.add_paragraph()
                pp.add_run('Page: ').bold = True
                pp.add_run(str(page))
                
                # Build context paragraph
                ctx = doc.add_paragraph()
                
                # Add previous sentence (normal)
                if prev_s:
                    ctx.add_run(prev_s + ' ')
                
                # Add matched sentence with yellow highlight
                # Find keyword within matched sentence and make it bold
                if kw and match_s:
                    pattern = regex.compile(regex.escape(kw), regex.IGNORECASE)
                    last_end = 0
                    for match in pattern.finditer(match_s):
                        # Add text before keyword (yellow, not bold)
                        if match.start() > last_end:
                            run_before = ctx.add_run(match_s[last_end:match.start()])
                            run_before.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        # Add keyword (yellow + bold)
                        keyword_run = ctx.add_run(match.group())
                        keyword_run.bold = True
                        keyword_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        last_end = match.end()
                    # Add remaining text of matched sentence (yellow, not bold)
                    if last_end < len(match_s):
                        run_after = ctx.add_run(match_s[last_end:])
                        run_after.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif match_s:
                    # No keyword, just highlight whole sentence
                    run_match = ctx.add_run(match_s)
                    run_match.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                # Add next sentence (normal)
                if next_s:
                    ctx.add_run(' ' + next_s)
                
                doc.add_paragraph('-' * 60)
            
            out_path = export_dir / f'{filename}.docx'
            doc.save(str(out_path))
            return send_file(str(out_path), as_attachment=True, download_name=f'{filename}.docx')
        
        return jsonify({'error': 'Invalid format'}), 400
    
    # -----------------------------------------------------------------------
    # WebSocket events – PDF Processing
    # -----------------------------------------------------------------------
    
    @socketio.on('start_processing')
    def handle_start_processing(data):
        """Start PDF processing in a background thread."""
        if state['processing']:
            emit('log_message', {'message': 'Processing is already running.'})
            return
        
        directory = data.get('directory', '')
        if not directory:
            emit('log_message', {'message': 'Please select a directory.'})
            return
        
        dir_path = Path(directory)
        if not dir_path.exists() or not dir_path.is_dir():
            emit('log_message', {'message': f'Directory not found: {directory}'})
            return
        
        pdf_files = list(dir_path.glob('*.pdf'))
        if not pdf_files:
            emit('log_message', {'message': f'No PDF files found in {directory}'})
            return
        
        # Reset state
        state['processing'] = True
        state['stop_flag'] = False
        state['log_messages'] = []
        state['file_statuses'] = []
        state['process_start_time'] = time.time()
        state['process_end_time'] = None
        state['last_completed_stats'] = None
        state['process_directory'] = str(dir_path)
        state['process_total_files'] = len(pdf_files)
        state['process_progress'] = 0
        
        # Parse options
        options = {
            'use_ocr': data.get('use_ocr', False),
            'create_references': data.get('create_references', False),
            'create_backups': data.get('create_backups', True),
            'move_problematic': data.get('move_problematic', True),
            'unnamed_dir': data.get('unnamed_dir', ''),
            'separate_ai_folder': data.get('separate_ai_folder', False),
            'categorize_options': {
                'by_journal': data.get('by_journal', False),
                'by_author': data.get('by_author', False),
                'by_year': data.get('by_year', False),
                'by_subject': data.get('by_subject', False),
            }
        }
        
        def run_processing():
            logger = logging.getLogger(f'litorganizer.web_worker_{id(threading.current_thread())}')
            logger.setLevel(logging.DEBUG)
            handler = SocketIOLogHandler()
            handler.setFormatter(logging.Formatter('%(message)s'))
            logger.addHandler(handler)
            
            try:
                unnamed_dir = options['unnamed_dir'] or None
                # Load API config for Gemini and other AI features
                api_config = load_api_config()
                
                processor = PDFProcessor(
                    directory=dir_path,
                    use_ocr=options['use_ocr'],
                    create_references=options['create_references'],
                    create_backups=options['create_backups'],
                    move_problematic=options['move_problematic'],
                    problematic_dir=unnamed_dir,
                    auto_analyze=False,
                    categorize_options=options['categorize_options'],
                    max_workers=4,
                    logger=logger,
                    api_config=api_config,
                    separate_ai_folder=options['separate_ai_folder'],
                )
                
                # Set event callback for Gemini AI status updates
                def gemini_event_callback(event_name, data):
                    try:
                        socketio.emit(event_name, data)
                    except Exception:
                        pass
                
                processor.event_callback = gemini_event_callback
                
                total_files = len(pdf_files)
                current_file = [0]
                
                original_process_file = processor.process_file
                
                def custom_process_file(file_path):
                    if state['stop_flag']:
                        return False
                    result = original_process_file(file_path)
                    
                    current_file[0] += 1
                    pct = int((current_file[0] / total_files) * 100)
                    
                    status = {
                        'filename': file_path.name,
                        'success': result,
                    }
                    state['file_statuses'].append(status)
                    state['process_progress'] = pct
                    
                    socketio.emit('file_processed', status)
                    socketio.emit('progress_update', {'percentage': pct})
                    return result
                
                processor.process_file = custom_process_file
                processor.process_files()
                
                state['process_end_time'] = time.time()
                
                # Build statistics
                stats = _build_stats(
                    processor, state['process_start_time'], state['process_end_time'],
                    options['categorize_options'], state['log_messages']
                )
                state['last_stats'] = stats
                
                # Store completed stats for page reload persistence
                completed_data = {
                    'processed': processor.processed_count,
                    'renamed': processor.renamed_count,
                    'problematic': processor.problematic_count,
                    'stats': stats,
                    'directory': state.get('process_directory', ''),
                }
                state['last_completed_stats'] = completed_data
                
                socketio.emit('processing_complete', completed_data)
            except Exception as e:
                socketio.emit('log_message', {'message': f'Error: {str(e)}'})
                logger.error(traceback.format_exc())
            finally:
                state['processing'] = False
                logger.removeHandler(handler)
        
        t = threading.Thread(target=run_processing, daemon=True)
        state['worker_thread'] = t
        
        socketio.emit('log_message', {'message': f'Found {len(pdf_files)} PDF files. Starting processing...'})
        socketio.emit('processing_started', {'total': len(pdf_files)})
        t.start()
    
    @socketio.on('stop_processing')
    def handle_stop_processing():
        """Signal the processing thread to stop."""
        state['stop_flag'] = True
        socketio.emit('log_message', {'message': 'Stopping processing...'})
    
    # -----------------------------------------------------------------------
    # WebSocket events – Keyword Search
    # -----------------------------------------------------------------------
    
    @socketio.on('start_search')
    def handle_start_search(data):
        """Start keyword search in a background thread."""
        if state['searching']:
            emit('log_message', {'message': 'Search is already running.'})
            return
        
        directory = data.get('directory', '')
        keyword = data.get('keyword', '').strip()
        
        if not directory:
            emit('log_message', {'message': 'Please select a directory.'})
            return
        if not keyword:
            emit('log_message', {'message': 'Please enter a keyword.'})
            return
        
        dir_path = Path(directory)
        if not dir_path.exists():
            emit('log_message', {'message': f'Directory not found: {directory}'})
            return
        
        exact_match = data.get('exact_match', False)
        case_sensitive = data.get('case_sensitive', False)
        use_regex = data.get('use_regex', False)
        
        if use_regex:
            try:
                re.compile(keyword)
            except re.error:
                emit('log_message', {'message': 'Invalid regular expression.'})
                return
        
        state['searching'] = True
        state['search_stop_flag'] = False
        state['search_results'] = []
        
        def run_search():
            import fitz  # PyMuPDF
            
            logger = logging.getLogger(f'litorganizer.web_search_{id(threading.current_thread())}')
            logger.setLevel(logging.DEBUG)
            handler = SocketIOLogHandler()
            handler.setFormatter(logging.Formatter('%(message)s'))
            logger.addHandler(handler)
            
            found_matches = 0
            processed_files = 0
            
            try:
                pdf_files = list(dir_path.glob('**/*.pdf'))
                total = len(pdf_files)
                
                if total == 0:
                    socketio.emit('log_message', {'message': f'No PDF files found in {directory}'})
                    socketio.emit('search_complete', {'processed': 0, 'found': 0})
                    return
                
                socketio.emit('log_message', {'message': f'Found {total} PDF files. Starting search...'})
                socketio.emit('search_started', {'total': total})
                
                # Prepare pattern
                if use_regex:
                    keyword_pattern = keyword
                else:
                    keyword_pattern = r'\b' + re.escape(keyword) + r'\b' if exact_match else re.escape(keyword)
                flags = 0 if case_sensitive else re.IGNORECASE
                
                def clean_text(text):
                    return re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
                
                def split_sentences(text):
                    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|!)\s', text)
                    return [s.strip() for s in sentences if s.strip()]
                
                def process_pdf(file_path):
                    nonlocal found_matches, processed_files
                    if state['search_stop_flag']:
                        return
                    
                    fname = file_path.name
                    
                    try:
                        # Extract DOI
                        doi = None
                        try:
                            doi = extract_doi(file_path, False)
                            if doi:
                                doi_start = doi.find('10.')
                                if doi_start >= 0:
                                    doi = doi[doi_start:]
                        except Exception:
                            pass
                        
                        # Extract text
                        with fitz.open(str(file_path)) as doc:
                            text_list = [page.get_text('text') for page in doc]
                        
                        for page_num, text in enumerate(text_list, start=1):
                            if state['search_stop_flag']:
                                return
                            if not text:
                                continue
                            text = clean_text(text)
                            sentences = split_sentences(text)
                            
                            for i, sentence in enumerate(sentences):
                                if state['search_stop_flag']:
                                    return
                                try:
                                    if re.search(keyword_pattern, sentence, flags):
                                        prev_s = ' '.join(sentences[max(i-1, 0):i]) if i > 0 else ''
                                        next_s = ' '.join(sentences[i+1:min(i+2, len(sentences))]) if i+1 < len(sentences) else ''
                                        
                                        result_row = [doi or '', fname, page_num, keyword, prev_s, sentence, next_s]
                                        state['search_results'].append(result_row)
                                        found_matches += 1
                                        
                                        socketio.emit('search_result', {
                                            'doi': doi or '',
                                            'filename': fname,
                                            'page': page_num,
                                            'keyword': keyword,
                                            'prev_sentence': prev_s,
                                            'matched_sentence': sentence,
                                            'next_sentence': next_s,
                                        })
                                except Exception:
                                    continue
                        
                        socketio.emit('search_file_processed', {'filename': fname, 'success': True})
                    except Exception as e:
                        socketio.emit('search_file_processed', {'filename': fname, 'success': False})
                        logger.error(f'Error processing {fname}: {e}')
                    finally:
                        processed_files += 1
                        pct = int(100 * processed_files / total)
                        state['search_progress'] = pct
                        socketio.emit('search_progress', {'percentage': pct})
                
                max_workers = min(os.cpu_count() or 4, 4)
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = {executor.submit(process_pdf, f): f for f in pdf_files}
                    for future in as_completed(futures):
                        if state['search_stop_flag']:
                            executor.shutdown(wait=False, cancel_futures=True)
                            break
                
                gc.collect()
                
                if found_matches > 0:
                    socketio.emit('log_message', {'message': f'Search completed. Found {found_matches} matches in {processed_files} files.'})
                else:
                    socketio.emit('log_message', {'message': f'Search completed. No matches found in {processed_files} files.'})
                
                socketio.emit('search_complete', {'processed': processed_files, 'found': found_matches})
                
            except Exception as e:
                socketio.emit('log_message', {'message': f'Search error: {str(e)}'})
                logger.error(traceback.format_exc())
                socketio.emit('search_complete', {'processed': processed_files, 'found': 0})
            finally:
                state['searching'] = False
                logger.removeHandler(handler)
        
        t = threading.Thread(target=run_search, daemon=True)
        state['search_thread'] = t
        t.start()
    
    @socketio.on('stop_search')
    def handle_stop_search():
        """Signal the search thread to stop."""
        state['search_stop_flag'] = True
        socketio.emit('log_message', {'message': 'Stopping search...'})
    
    # -----------------------------------------------------------------------
    # Statistics helper
    # -----------------------------------------------------------------------
    
    def _build_stats(processor, start_time, end_time, categorize_options, log_messages):
        """Build a statistics dictionary from processing results."""
        processed = processor.processed_count
        renamed = processor.renamed_count
        problematic = processor.problematic_count
        success_rate = round((renamed / processed) * 100, 1) if processed > 0 else 0
        
        total_time = round(end_time - start_time, 2) if start_time and end_time else 0
        time_per_file = round(total_time / processed, 2) if processed > 0 else 0
        speed = round(processed / total_time, 2) if total_time > 0 else 0
        estimated_memory = round(processed * 0.5, 1)
        manual_time = processed * 30
        time_saved = round((manual_time - total_time) / 60, 1) if total_time > 0 else 0
        
        doi_detection = success_rate
        metadata_quality = round(success_rate * 0.9, 1)
        
        cat_enabled = sum([
            categorize_options.get('by_journal', False),
            categorize_options.get('by_author', False),
            categorize_options.get('by_year', False),
            categorize_options.get('by_subject', False),
        ])
        categorization_quality = round(100 * (cat_enabled / 4), 1) if cat_enabled > 0 else 0
        
        # API source analysis
        log_text = '\n'.join(log_messages)
        api_sources = {}
        api_matches = re.findall(r'Sufficient metadata found for .*? via (\w+)', log_text)
        for api in api_matches:
            api_sources[api] = api_sources.get(api, 0) + 1
        
        # Error breakdown
        errors = {}
        if problematic > 0:
            error_patterns = {
                'Missing DOI': r'No DOI found in',
                'Insufficient Metadata': r'Insufficient or no metadata found for DOI',
                'PDF Read Error': r'PDF Processing Error \(read\)',
                'PDF Encrypted': r'PDF Processing Error \(encrypted\)',
                'DOI Extraction Error': r'Error extracting DOI from',
                'API Error': r'API Error \(network/http\)',
                'Metadata Fetch Error': r'Error fetching metadata for DOI',
                'File System Error': r'File System Error',
                'Rename/Move Error': r'Error renaming/moving file',
                'Categorization Error': r'Error during file categorization attempt|Error calling categorize_file',
                'Unexpected Error': r'Unexpected Error processing file',
            }
            counted = 0
            for label, pattern in error_patterns.items():
                count = len(re.findall(pattern, log_text, re.IGNORECASE))
                if count > 0:
                    errors[label] = count
                    counted += count
            other = max(0, problematic - counted)
            if other > 0:
                errors['Other'] = other
        
        # Category stats
        category_counts = processor.category_counts
        categorized_file_count = processor.categorized_file_count
        
        return {
            'processed': processed,
            'renamed': renamed,
            'problematic': problematic,
            'success_rate': success_rate,
            'total_time': total_time,
            'time_per_file': time_per_file,
            'speed': speed,
            'estimated_memory': estimated_memory,
            'time_saved': time_saved,
            'doi_detection': doi_detection,
            'metadata_quality': metadata_quality,
            'categorization_quality': categorization_quality,
            'api_sources': api_sources,
            'errors': errors,
            'category_counts': category_counts,
            'categorized_file_count': categorized_file_count,
            'categorize_options': categorize_options,
        }
    
    return app, socketio


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def launch_web(logger: Optional[logging.Logger] = None, port: int = 5000):
    """
    Launch the web-based interface.
    
    Starts a Flask-SocketIO server on localhost and opens the default browser.
    
    Args:
        logger: Optional logger instance
        port: Port number (default 5000)
    """
    app_root = Path(__file__).resolve().parent.parent.parent
    os.chdir(app_root)
    
    # Setup logging
    log_dir = app_root / 'logs'
    log_dir.mkdir(exist_ok=True)
    
    current_time = datetime.now().strftime('%Y%m%d_%H%M%S')
    log_file = log_dir / f'web_{current_time}.log'
    
    if logger is None:
        logger = logging.getLogger('litorganizer')
    
    logger.setLevel(logging.DEBUG)
    
    file_handler = logging.FileHandler(log_file, 'w', 'utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
    logger.addHandler(file_handler)
    
    logger.info('Starting LitOrganizer web interface')
    logger.info(f'Log file: {log_file}')
    
    app, socketio = create_app()
    
    # Open browser after a short delay
    def open_browser():
        import time as _time
        _time.sleep(1.5)
        webbrowser.open(f'http://localhost:{port}')
    
    threading.Thread(target=open_browser, daemon=True).start()
    
    print(f'\n  LitOrganizer Web Interface')
    print(f'  Running at: http://localhost:{port}')
    print(f'  Press Ctrl+C to stop\n')
    
    socketio.run(app, host='127.0.0.1', port=port, debug=False, allow_unsafe_werkzeug=True)
