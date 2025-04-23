"""
PyQt5-based graphical user interface for the PDF Citation Tool.
"""

import os
import sys
import time
import logging
import traceback
import subprocess
import platform
from typing import Optional, List, Dict, Any, Union, Tuple
from pathlib import Path
from datetime import datetime
import re

# Platform-specific imports
if platform.system() == "Windows":
    from ctypes import windll, byref, c_int, sizeof
else:
    # Dummy objects for non-Windows platforms
    windll = None
    byref = None
    c_int = None
    sizeof = None

try:
    from PyQt5.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QPushButton, QLabel, QLineEdit, QFileDialog, QCheckBox,
        QTextEdit, QProgressBar, QMessageBox, QGroupBox, QGridLayout,
        QFrame, QSplitter, QTableWidget, QTableWidgetItem, QHeaderView,
        QTabWidget, QListWidget, QListWidgetItem, QStyleFactory,
        QComboBox, QScrollArea, QSizePolicy, QGraphicsDropShadowEffect,
        QFormLayout, QDialog
    )
    from PyQt5.QtCore import (
        QThread, pyqtSignal, Qt, QTimer, QSize, QSettings
    )
    from PyQt5.QtGui import (
        QColor, QFont, QPixmap, QIcon, QTextCursor, QPainter
    )
    QApplication.setStyle(QStyleFactory.create('Fusion'))
except ImportError as e:
    print(f"PyQt5 is required but not installed: {e}")
    sys.exit(1)

from modules.core import PDFProcessor
from modules.utils.file_utils import ensure_dir, get_version
from modules.utils.pdf_metadata_extractor import load_api_config


class WorkerThread(QThread):
    """
    Worker thread for processing PDF files in the background.
    """
    
    # Signals for progress reporting
    progress_update = pyqtSignal(str)
    file_processed = pyqtSignal(str, bool)
    processing_complete = pyqtSignal(int, int, int)
    error_occurred = pyqtSignal(str)
    progress_percentage = pyqtSignal(int)  # Signal for percentage-based progress
    
    def __init__(
        self,
        directory: Path,
        use_ocr: bool = False,
        create_references: bool = False,
        create_backups: bool = True,
        move_problematic: bool = False,
        problematic_dir: Optional[str] = None,
        auto_analyze: bool = False,
        categorize_options: Dict[str, bool] = None,
        logger: Optional[logging.Logger] = None
    ):
        """
        Initialize the worker thread.
        
        Args:
            directory (Path): Path to the PDF directory
            use_ocr (bool): Whether to use OCR for scanned PDFs
            create_references (bool): Whether to create a references file
            create_backups (bool): Whether to create backups of original files
            move_problematic (bool): Whether to move problematic files to a separate folder
            problematic_dir (Optional[str]): Path to the problematic files directory
            auto_analyze (bool): Whether to auto-analyze PDFs
            categorize_options (Dict[str, bool]): Categorization options
            logger (Optional[logging.Logger]): Logger instance
        """
        super().__init__()
        self.directory = directory
        self.use_ocr = use_ocr
        self.create_references = create_references
        self.create_backups = create_backups
        self.move_problematic = move_problematic
        self.problematic_dir = problematic_dir
        self.auto_analyze = False
        self.categorize_options = categorize_options or {}
        self.logger = logger
        self.terminate_flag = False
        
        # File counters
        self.total_files = 0
        self.current_file = 0
    
    def run(self):
        """
        Run the worker thread.
        """
        try:
            # Create custom logger for this thread that will emit signals
            thread_logger = logging.getLogger(f'litorganizer.worker_{id(self)}')
            thread_logger.setLevel(self.logger.level)
            
            # Set our custom handler that emits signals
            signal_handler = SignalLogHandler(self.progress_update)
            thread_logger.addHandler(signal_handler)
            
            # Create processor instance
            processor = PDFProcessor(
                directory=self.directory,
                use_ocr=self.use_ocr,
                create_references=self.create_references,
                create_backups=self.create_backups,
                move_problematic=self.move_problematic,
                problematic_dir=self.problematic_dir,
                auto_analyze=self.auto_analyze,
                categorize_options=self.categorize_options,
                max_workers=4,
                logger=thread_logger
            )
            
            # Get all PDF files and set total
            pdf_files = list(self.directory.glob("*.pdf"))
            self.total_files = len(pdf_files)
            thread_logger.info(f"Found {self.total_files} PDF files to process")
            
            # Create a custom process_file method that emits signals
            original_process_file = processor.process_file
            
            def custom_process_file(file_path):
                if self.terminate_flag:
                    return False
                
                result = original_process_file(file_path)
                self.file_processed.emit(file_path.name, result)
                
                # Calculate percentage and send signal
                self.current_file += 1
                percentage = int((self.current_file / self.total_files) * 100)
                self.progress_percentage.emit(percentage)
                
                return result
            
            # Replace the method
            processor.process_file = custom_process_file
            
            # Start processing
            processor.process_files()
            
            # Emit completion signal
            self.processing_complete.emit(
                processor.processed_count,
                processor.renamed_count,
                processor.problematic_count
            )
        
        except Exception as e:
            self.error_occurred.emit(str(e))
            self.logger.error(f"Error in worker thread: {e}")
    
    def terminate(self):
        """
        Terminate the worker thread.
        """
        self.terminate_flag = True
        self.wait()


class SignalLogHandler(logging.Handler):
    """
    Custom log handler that emits signals instead of writing to a file or console.
    """
    
    def __init__(self, signal):
        """
        Initialize the handler with a signal to emit log records.
        
        Args:
            signal: Signal to emit log records
        """
        super().__init__()
        self.signal = signal
    
    def emit(self, record):
        """
        Emit a log record as a signal.
        
        Args:
            record: Log record to emit
        """
        log_entry = self.format(record)
        self.signal.emit(log_entry)


class MainWindow(QMainWindow):
    """
    Main application window.
    """
    
    def __init__(self, logger: logging.Logger):
        """
        Initialize the main window.
        
        Args:
            logger (logging.Logger): Logger instance
        """
        super().__init__()
        self.logger = logger
        self.setWindowTitle("LitOrganizer - Organize your academic literature efficiently")
        self.setWindowIcon(QIcon("resources/icon_windows.png"))  # Changed back to icon_windows.png
        self.resize(1450, 1000)  # Larger window size for better visibility
        self.setMinimumSize(1350, 950)  # Larger minimum size
        
        # Set window theme to light/dark (Windows-specific)
        if platform.system() == "Windows" and windll is not None:
            try:
                # Windows 11 theme styling - Use light theme (0) instead of dark theme (1)
                windll.dwmapi.DwmSetWindowAttribute(
                    int(self.winId()), 19, byref(c_int(0)), sizeof(c_int)
                )
            except Exception:
                pass  # May not work on Windows versions below 10, silently continue
        
        # Default settings
        self.directory = ""
        self.unnamed_article_dir = ""
        self.categorize_by_journal = False
        self.categorize_by_author = False
        self.categorize_by_year = False
        self.categorize_by_subject = False
        
        # Performance tracking
        self.process_start_time = None
        self.process_end_time = None
        
        # Setup UI components
        self.setup_ui()
        
        # Load saved settings
        self.load_settings()
        
        # Worker thread
        self.worker_thread = None
        
        self.logger.info("GUI initialized")
        
        # Initialize statistics graphs
        self.initialize_statistics()
    
    def initialize_statistics(self):
        """
        Initialize statistics displays with empty data.
        """
        # Reset performance metrics
        if hasattr(self, 'stat_total_time'):
            self.stat_total_time.setText("Total Processing Time: 0 seconds")
            self.stat_per_file_time.setText("Average Time Per File: 0 seconds")
            self.stat_processing_speed.setText("Processing Speed: 0 files/second")
            self.stat_memory_usage.setText("Estimated Memory Usage: 0 MB")
            
            # Reset performance progress bars
            if hasattr(self, 'total_time_progress'):
                self.total_time_progress.setValue(0)
                self.per_file_time_progress.setValue(0)
                self.processing_speed_progress.setValue(0)
                self.memory_usage_progress.setValue(0)
        
        # Reset accuracy metrics
        if hasattr(self, 'stat_doi_detection'):
            self.stat_doi_detection.setText("DOI Detection Rate: 0%")
            self.stat_metadata_quality.setText("Metadata Quality: 0%")
            self.stat_categorization_quality.setText("Categorization Quality: 0%")
            self.stat_time_savings.setText("Estimated Time Saved: 0 minutes")
            
            # Reset accuracy progress bars
            if hasattr(self, 'doi_detection_progress'):
                self.doi_detection_progress.setValue(0)
                self.metadata_quality_progress.setValue(0)
                self.categorization_quality_progress.setValue(0)
                self.time_savings_progress.setValue(0)
            
        # Reset top authors and journals
        if hasattr(self, 'top_authors_list'):
            self.top_authors_list.clear()
            self.top_authors_list.addItem("No data available yet")
            
        if hasattr(self, 'top_journals_list'):
            self.top_journals_list.clear()
            self.top_journals_list.addItem("No data available yet")
            
        if hasattr(self, 'year_distribution_list'):
            self.year_distribution_list.clear()
            self.year_distribution_list.addItem("No data available yet")
            
        if hasattr(self, 'subject_list'):
            self.subject_list.clear()
            self.subject_list.addItem("No subject data available yet")
            
        if hasattr(self, 'api_stats_list'):
            self.api_stats_list.clear()
            self.api_stats_list.addItem("No API source information found")
            
        if hasattr(self, 'error_breakdown_list'):
            self.error_breakdown_list.clear()
            self.error_breakdown_list.addItem("No errors identified yet")
    
    def setup_ui(self):
        """
        Set up the user interface.
        """
        # Application icons
        save_icon = QIcon("resources/save.png")
        
        if save_icon.isNull():
            # If icon failed to load, create emoji icon
            save_pixmap = QPixmap(24, 24)
            save_pixmap.fill(Qt.transparent)
            painter = QPainter(save_pixmap)
            painter.setFont(QFont("Segoe UI Emoji", 14))
            painter.drawText(save_pixmap.rect(), Qt.AlignCenter, "💾")
            painter.end()
            save_icon = QIcon(save_pixmap)
        
        # Windows 11 color palette - modern palette for better visibility
        primary_color = "#0078D4"     # Windows blue (darker tone)
        accent_color = "#50ABF1"      # Light blue (accent - harmonized)
        success_color = "#107C10"     # Green (more lively)
        warning_color = "#F8A800"     # Yellow (harmonized darker tone)
        error_color = "#D83B01"       # Red (harmonized darker tone)
        text_color = "#202020"        # Main text
        secondary_text = "#5F6368"    # Secondary text
        divider_color = "#E1E1E1"     # Dividers
        
        background_color = "#F9F9F9"  # Background
        card_color = "#FFFFFF"        # Card background
        hover_color = "#F2F2F2"       # Hover background
        subtle_hover = "#F5F5F5"      # Subtle hover
        
        # QT's consistent UI settings
        self.groupbox_style = "QGroupBox { font-weight: bold; }"  # Common style for all groupboxes
        border_radius = "8px"         # More consistent rounded corners
        input_radius = "6px"          # More consistent input corners
        
        # Set application styles
        self.setStyleSheet(f"""
            QMainWindow {{
                background-color: {card_color};
            }}
        """)
        
        # Main layout
        central_widget = QWidget()
        central_widget.setStyleSheet(f"""
            QWidget {{
                background-color: {background_color};
                color: {text_color};
                font-family: 'Segoe UI', sans-serif;
                font-size: 9.5pt;
            }}
            QGroupBox {{
                border: 1px solid {divider_color};
                border-radius: {border_radius};
                margin-top: 1.5ex;
                padding: 10px;
                background-color: {card_color};
                font-weight: normal;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 0 8px;
                color: {primary_color};
                font-weight: 600;
                font-size: 10pt;
            }}
            QLineEdit {{
                border: 1px solid {divider_color};
                border-radius: {input_radius};
                padding: 8px;
                background-color: {card_color};
                selection-background-color: {primary_color};
                min-height: 16px;
                font-size: 9pt;
            }}
            QLineEdit:focus {{
                border-color: {primary_color};
                background-color: {card_color};
            }}
            QPushButton {{
                border-radius: {input_radius};
                padding: 8px 14px;
                background-color: {primary_color};
                color: white;
                border: none;
                font-weight: 600;
                min-height: 16px;
                font-size: 9pt;
            }}
            QPushButton:hover {{
                background-color: #006CBE;
            }}
            QPushButton:pressed {{
                background-color: #005AA0;
            }}
            QPushButton:disabled {{
                background-color: {divider_color};
                color: {secondary_text};
            }}
            QCheckBox {{
                spacing: 8px;
                font-size: 9pt;
            }}
            QCheckBox::indicator {{
                width: 20px;
                height: 20px;
                border-radius: 3px;
            }}
            QTabWidget::pane {{
                border: 1px solid {divider_color};
                border-radius: {border_radius};
                background-color: {background_color};
                top: -1px;
            }}
            QTabBar::tab {{
                background-color: transparent;
                border: none;
                border-bottom: 2px solid transparent;
                padding: 8px 20px;
                margin-right: 6px;
                font-weight: normal;
                color: {secondary_text};
                font-size: 9.5pt;
                font-family: "Segoe UI", Verdana, sans-serif;
                min-width: 100px;
            }}
            QTabBar::tab:selected {{
                background-color: transparent;
                border-bottom: 2px solid {primary_color};
                color: {primary_color};
                font-weight: 600;
            }}
            QTabBar::tab:hover:!selected {{
                background-color: transparent;
                font-weight: bold;
                font-size: 10pt;
            }}
            QScrollBar:vertical {{
                border: none;
                background-color: {hover_color};
                width: 12px;
                margin: 0px;
                border-radius: 6px;
            }}
            QScrollBar::handle:vertical {{
                background-color: #C0C0C0;
                min-height: 28px;
                border-radius: 6px;
            }}
            QScrollBar::handle:vertical:hover {{
                background-color: #A0A0A0;
            }}
            QScrollBar:horizontal {{
                border: none;
                background-color: {hover_color};
                height: 12px;
                margin: 0px;
                border-radius: 6px;
            }}
            QScrollBar::handle:horizontal {{
                background-color: #C0C0C0;
                min-width: 28px;
                border-radius: 6px;
            }}
            QScrollBar::handle:horizontal:hover {{
                background-color: #A0A0A0;
            }}
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal,
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal {{
                width: 0px;
                background: none;
            }}
            QProgressBar {{
                border: none;
                border-radius: 6px;
                text-align: center;
                background-color: {divider_color};
                height: 12px;
                font-size: 8pt;
                font-weight: bold;
            }}
            QProgressBar::chunk {{
                background-color: {primary_color};
                border-radius: 6px;
            }}
            QLabel {{
                font-size: 9.5pt;
            }}
            QTextEdit {{
                border: 1px solid {divider_color};
                border-radius: {input_radius};
                padding: 4px;
                background-color: {card_color};
                selection-background-color: {primary_color};
                font-size: 9pt;
            }}
        """)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(10)
        self.setCentralWidget(central_widget)
        
        # Add header with logo only
        header_layout = QHBoxLayout()
        header_layout.setContentsMargins(0, 0, 0, 20)  # Add some space at the bottom
        
        # Logo
        logo_label = QLabel()
        logo_pixmap = QPixmap("resources/logo.png")
        if not logo_pixmap.isNull():
            logo_pixmap = logo_pixmap.scaled(250, 250, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            logo_label.setPixmap(logo_pixmap)
        else:
            # Fallback if logo.png doesn't exist - use a document emoji as a placeholder
            logo_label.setText("📑")
            logo_label.setFont(QFont("Segoe UI", 24))
            logo_label.setStyleSheet(f"color: {primary_color};")
        
        logo_label.setAlignment(Qt.AlignCenter)
        header_layout.addWidget(logo_label)
        
        # Add empty space in the middle
        header_layout.addStretch(1)
        
        # BibexPy logo on the right (for branding)
        bibexpy_layout = QVBoxLayout()
        
        # BibexPy icon
        bibexpy_icon_label = QLabel()
        bibexpy_icon = QPixmap("resources/logo_bibewxpy.webp")
        if not bibexpy_icon.isNull():
            bibexpy_icon = bibexpy_icon.scaled(180, 60, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            bibexpy_icon_label.setPixmap(bibexpy_icon)
            bibexpy_icon_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        
        # BibexPy slogan
        bibexpy_slogan = QLabel("Bibliometrics Experience with Python")
        bibexpy_slogan.setFont(QFont("Segoe UI", 10))
        bibexpy_slogan.setStyleSheet(f"color: {primary_color};")
        bibexpy_slogan.setAlignment(Qt.AlignRight)
        
        bibexpy_website = QLabel("bibexpy.com")
        bibexpy_website.setFont(QFont("Segoe UI", 9))
        bibexpy_website.setStyleSheet(f"color: {secondary_text};")
        bibexpy_website.setAlignment(Qt.AlignRight)

        bibexpy_layout.addWidget(bibexpy_icon_label)
        bibexpy_layout.addWidget(bibexpy_slogan)
        bibexpy_layout.addWidget(bibexpy_website)
        bibexpy_layout.setSpacing(2)
        
        header_layout.addLayout(bibexpy_layout)
        
        # Add the header
        main_layout.addLayout(header_layout)
        
        # Create tabs
        self.tabs = QTabWidget()
        self.tabs.setDocumentMode(True)  # More modern tab view
        self.tabs.setElideMode(Qt.ElideRight)  # Clip mode for long tab names
        main_layout.addWidget(self.tabs)
        
        # Main tab
        main_tab = QWidget()
        main_tab_layout = QVBoxLayout(main_tab)
        main_tab_layout.setSpacing(5)  # Reduce vertical spacing
        main_tab_layout.setContentsMargins(5, 5, 5, 5)  # Reduce margin spacing
        
        # Create a widget for the upper section to contain all upper components
        top_container = QWidget()
        top_layout = QVBoxLayout(top_container)
        top_layout.setContentsMargins(0, 0, 0, 0)  # Remove margins completely
        top_layout.setSpacing(5)  # Reduce spacing
        
        # PDF Directory selection
        dir_group = QGroupBox("PDF Directory")
        dir_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(dir_group)
        dir_layout = QVBoxLayout(dir_group)
        dir_layout.setContentsMargins(6, 12, 6, 6)  # Reduce margins
        dir_layout.setSpacing(3)  # Reduce internal spacing
        
        dir_input_layout = QHBoxLayout()
        self.dir_edit = QLineEdit()
        self.dir_edit.setPlaceholderText("Select a directory containing PDF files")
        dir_input_layout.addWidget(self.dir_edit)
        
        # Browse buttons with different styles
        browse_button_style = f"""
            QPushButton {{
                background-color: {hover_color};
                border: 1px solid {divider_color};
                border-radius: 8px;
                padding: 9px 15px;
                color: {text_color};
                font-weight: normal;
            }}
            QPushButton:hover {{
                background-color: {subtle_hover};
                border-color: {primary_color};
            }}
            QPushButton:pressed {{
                background-color: {divider_color};
            }}
        """
        
        self.browse_btn = QPushButton("Browse...")
        self.browse_btn.clicked.connect(self.browse_directory)
        self.browse_btn.setStyleSheet(browse_button_style)
        dir_input_layout.addWidget(self.browse_btn)
        
        self.open_dir_btn = QPushButton("📁")
        self.open_dir_btn.setToolTip("Open the selected PDF directory.")
        self.open_dir_btn.clicked.connect(lambda: self.open_folder(self.dir_edit.text()))
        self.open_dir_btn.setStyleSheet(browse_button_style)
        dir_input_layout.addWidget(self.open_dir_btn)
        
        dir_layout.addLayout(dir_input_layout)
        top_layout.addWidget(dir_group)
        
        # Unnamed Article Files Directory
        unclassified_group = QGroupBox("Unnamed Article Files")
        unclassified_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(unclassified_group)
        unclassified_layout = QVBoxLayout(unclassified_group)
        unclassified_layout.setContentsMargins(6, 12, 6, 6)  # Reduce margins
        unclassified_layout.setSpacing(3)  # Reduce internal spacing
        
        unclassified_input_layout = QHBoxLayout()
        self.unclassified_dir_edit = QLineEdit()
        self.unclassified_dir_edit.setPlaceholderText("Directory for Unnamed Article files (creates 'Unnamed Article' subfolder if empty)")
        unclassified_input_layout.addWidget(self.unclassified_dir_edit)
        
        self.browse_unclassified_btn = QPushButton("Browse...")
        self.browse_unclassified_btn.clicked.connect(self.browse_unclassified_directory)
        self.browse_unclassified_btn.setStyleSheet(browse_button_style)
        unclassified_input_layout.addWidget(self.browse_unclassified_btn)
        
        self.open_unclassified_btn = QPushButton("📁")
        self.open_unclassified_btn.setToolTip("Open the Unnamed Article files directory.")
        self.open_unclassified_btn.clicked.connect(lambda: self.open_folder(self.unclassified_dir_edit.text()))
        self.open_unclassified_btn.setStyleSheet(browse_button_style)
        unclassified_input_layout.addWidget(self.open_unclassified_btn)
        
        unclassified_layout.addLayout(unclassified_input_layout)
        top_layout.addWidget(unclassified_group)
        
        # Create processing options hidden (not shown to user)
        # All processing options are active by default
        self.use_ocr_check = QCheckBox("Use OCR for text extraction (requires Tesseract)")
        self.use_ocr_check.setChecked(False)
        
        self.create_backups_check = QCheckBox("Create backups of original files")
        self.create_backups_check.setChecked(True)
        
        self.move_problematic_check = QCheckBox("Move Unnamed Article files to separate folder")
        self.move_problematic_check.setChecked(True)
        
        self.create_references_check = QCheckBox("Create references file")
        self.create_references_check.setChecked(True)
        
        # Categorization options
        categorize_group = QGroupBox("Categorization Options")
        categorize_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(categorize_group)
        categorize_layout = QVBoxLayout(categorize_group)
        categorize_layout.setContentsMargins(6, 12, 6, 6)  # Reduce margins
        categorize_layout.setSpacing(3)  # Reduce internal spacing
        
        categorize_checkbox_layout = QGridLayout()
        
        # Function to create styled checkboxes
        def create_styled_checkbox(text, tooltip):
            checkbox = QCheckBox(text)
            checkbox.setToolTip(tooltip)
            checkbox.setChecked(True)  # Active by default
            checkbox.setStyleSheet(f"""
                QCheckBox {{
                    padding: 8px;
                    font-size: 11pt;
                    color: {text_color};
                }}
                QCheckBox::indicator {{
                    width: 22px;
                    height: 22px;
                    border-radius: 4px;
                }}
                QCheckBox::indicator:unchecked {{
                    image: url(resources/unchecked.png);
                }}
                QCheckBox::indicator:checked {{
                    image: url(resources/check.png);
                }}
                QCheckBox::indicator:unchecked:hover {{
                    border: 1px solid {primary_color};
                }}
            """)
            return checkbox
        
        self.categorize_by_journal_check = create_styled_checkbox("By Journal", "Create folders organized by journal name")
        categorize_checkbox_layout.addWidget(self.categorize_by_journal_check, 0, 0)
        
        self.categorize_by_author_check = create_styled_checkbox("By Author (First Author)", "Create folders organized by author name")
        categorize_checkbox_layout.addWidget(self.categorize_by_author_check, 0, 1)
        
        self.categorize_by_year_check = create_styled_checkbox("By Year", "Create folders organized by publication year")
        categorize_checkbox_layout.addWidget(self.categorize_by_year_check, 1, 0)
        
        self.categorize_by_subject_check = create_styled_checkbox("By Subject", "Create folders organized by subject or category")
        categorize_checkbox_layout.addWidget(self.categorize_by_subject_check, 1, 1)
        
        categorize_layout.addLayout(categorize_checkbox_layout)
        top_layout.addWidget(categorize_group)
        
        # Buttons
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(20)  # Increase spacing between buttons
        
        self.start_btn = QPushButton("Start Processing")
        self.start_btn.clicked.connect(self.start_processing)
        self.start_btn.setMinimumSize(350, 50)  # Larger and wider button
        self.start_btn.setStyleSheet("""
            QPushButton {
                background-color: #107C10;
                color: white;
                border: none;
                padding: 14px 30px;
                border-radius: 10px;
                font-weight: 600;
                font-size: 14px;
                text-align: center;
            }
            QPushButton:hover {
                background-color: #0A6A0A;
            }
            QPushButton:pressed {
                background-color: #095909;
            }
        """)
        buttons_layout.addWidget(self.start_btn, 3)  # 3:1 ratio for width
        
        self.stop_btn = QPushButton("Stop")
        self.stop_btn.clicked.connect(self.stop_processing)
        self.stop_btn.setEnabled(False)
        self.stop_btn.setMinimumSize(120, 50)  # Smaller button
        self.stop_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {error_color};
                color: white;
                border: none;
                padding: 14px 20px;
                border-radius: 10px;
                font-weight: 600;
                font-size: 14px;
            }}
            QPushButton:hover {{
                background-color: #D31020;
            }}
            QPushButton:pressed {{
                background-color: #B80E1B;
            }}
            QPushButton:disabled {{
                background-color: {divider_color};
                color: {secondary_text};
            }}
        """)
        buttons_layout.addWidget(self.stop_btn, 1)  # 1 unit width
        
        # Clear buttons are removed
        # Buttons layout directly added to main layout
        top_layout.addLayout(buttons_layout)
        
        # Size policy for top container
        top_size_policy = QSizePolicy(QSizePolicy.Preferred, QSizePolicy.Maximum)
        top_size_policy.setVerticalStretch(0)
        top_container.setSizePolicy(top_size_policy)
        
        # Add top container to main tab
        main_tab_layout.addWidget(top_container, 0)  # 0 stretch factor
        
        # Main content area - New design
        content_container = QWidget()
        content_layout = QVBoxLayout(content_container)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)
        
        # Collapsible panels with QSplitter (side by side)
        horizontal_splitter = QSplitter(Qt.Horizontal)
        horizontal_splitter.setHandleWidth(8)
        horizontal_splitter.setChildrenCollapsible(False)
        horizontal_splitter.setOpaqueResize(True)
        horizontal_splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #F9F9F9;
                border: 1px solid #E1E1E1;
                border-radius: 2px;
            }
            QSplitter::handle:hover {
                background-color: #0078D4;
            }
        """)
        
        # RESULTS PANEL - LEFT
        results_panel = QWidget()
        results_panel.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        results_layout = QVBoxLayout(results_panel)
        results_layout.setContentsMargins(8, 8, 4, 8)
        
        # Results header - add collapse button
        results_header = QHBoxLayout()
        
        results_title = QLabel("Results")
        results_title.setStyleSheet(f"""
            font-weight: 600;
            font-size: 12px;
            color: {primary_color};
        """)
        results_header.addWidget(results_title)
        
        results_header.addStretch(1)
        
        # Save results button
        self.save_results_btn = QPushButton()
        self.save_results_btn.setToolTip("Save results to a file")
        self.save_results_btn.setFixedSize(34, 34)
        self.save_results_btn.setIcon(QIcon("resources/save.png"))
        self.save_results_btn.setIconSize(QSize(18, 18))
        self.save_results_btn.clicked.connect(self.save_results)
        self.save_results_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                border-radius: 17px;
                padding: 4px;
            }
            QPushButton:hover {
                background-color: #F2F2F2;
            }
            QPushButton:pressed {
                background-color: #E1E1E1;
            }
        """)
        results_header.addWidget(self.save_results_btn)
        
        # Clear results button
        self.clear_results_btn = QPushButton()
        self.clear_results_btn.setToolTip("Clear results")
        self.clear_results_btn.setFixedSize(34, 34)
        self.clear_results_btn.setIcon(QIcon("resources/dust.png"))
        self.clear_results_btn.setIconSize(QSize(18, 18))
        self.clear_results_btn.clicked.connect(self.clear_status_list)
        self.clear_results_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                border-radius: 17px;
                padding: 4px;
            }
            QPushButton:hover {
                background-color: #F2F2F2;
            }
            QPushButton:pressed {
                background-color: #E1E1E1;
            }
        """)
        results_header.addWidget(self.clear_results_btn)
        results_layout.addLayout(results_header)
        
        # Modern card design for results
        results_card = QFrame()
        results_card.setStyleSheet(f"""
            QFrame {{
                background-color: {card_color};
                border-radius: 10px;
                border: 1px solid {divider_color};
            }}
        """)
        results_card_layout = QVBoxLayout(results_card)
        results_card_layout.setContentsMargins(2, 2, 2, 2)
        
        self.results_list = QListWidget()
        self.results_list.setAlternatingRowColors(True)
        self.results_list.setStyleSheet(f"""
            QListWidget {{
                background-color: {card_color};
                border: none;
                padding: 4px;
                alternate-background-color: {subtle_hover};
            }}
            QListWidget::item {{
                padding: 8px;
                border-radius: 5px;
                margin: 2px 0px;
            }}
            QListWidget::item:selected {{
                background-color: {primary_color}30;
                color: {text_color};
                border-left: 3px solid {primary_color};
            }}
            QListWidget::item:hover:!selected {{
                background-color: {hover_color};
            }}
        """)
        self.add_shadow_effect(results_card, blur_radius=10, x_offset=0, y_offset=2)
        results_card_layout.addWidget(self.results_list)
        results_layout.addWidget(results_card)
        
        # LOG PANEL - RIGHT
        log_panel = QWidget()
        log_panel.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        log_layout = QVBoxLayout(log_panel)
        log_layout.setContentsMargins(4, 8, 8, 8)
        
        # Log header
        log_header = QHBoxLayout()
        
        log_title = QLabel("Processing Log")
        log_title.setStyleSheet(f"""
            font-weight: 600;
            font-size: 12px;
            color: {primary_color};
        """)
        log_header.addWidget(log_title)
        
        log_header.addStretch(1)
        
        # Clear log button
        self.clear_log_btn = QPushButton()
        self.clear_log_btn.setToolTip("Clear log")
        self.clear_log_btn.setFixedSize(34, 34)
        self.clear_log_btn.setIcon(QIcon("resources/dust.png"))
        self.clear_log_btn.setIconSize(QSize(18, 18))
        self.clear_log_btn.clicked.connect(self.clear_log)
        self.clear_log_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                border-radius: 17px;
                padding: 4px;
            }
            QPushButton:hover {
                background-color: #F2F2F2;
            }
            QPushButton:pressed {
                background-color: #E1E1E1;
            }
        """)
        log_header.addWidget(self.clear_log_btn)
        
        # Save log button
        self.save_log_btn = QPushButton()
        self.save_log_btn.setToolTip("Save log to a file")
        self.save_log_btn.setFixedSize(34, 34)
        self.save_log_btn.setIcon(QIcon("resources/save.png"))
        self.save_log_btn.setIconSize(QSize(18, 18))
        self.save_log_btn.clicked.connect(self.save_log)
        self.save_log_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                border-radius: 17px;
                padding: 4px;
            }
            QPushButton:hover {
                background-color: #F2F2F2;
            }
            QPushButton:pressed {
                background-color: #E1E1E1;
            }
        """)
        log_header.addWidget(self.save_log_btn)
        log_layout.addLayout(log_header)
        
        # Modern card design for log
        log_card = QFrame()
        log_card.setStyleSheet(f"""
            QFrame {{
                background-color: {card_color};
                border-radius: 10px;
                border: 1px solid {divider_color};
            }}
        """)
        log_card_layout = QVBoxLayout(log_card)
        log_card_layout.setContentsMargins(2, 2, 2, 2)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setFont(QFont("Cascadia Mono", 9))
        self.log_text.setStyleSheet(f"""
            QTextEdit {{
                background-color: {card_color};
                border: none;
                padding: 8px;
                selection-background-color: {primary_color}40;
                font-family: 'Cascadia Mono', 'Consolas', monospace;
            }}
        """)
        self.add_shadow_effect(log_card, blur_radius=10, x_offset=0, y_offset=2)
        log_card_layout.addWidget(self.log_text)
        log_layout.addWidget(log_card)
        
        # Add panels to splitter
        horizontal_splitter.addWidget(results_panel)
        horizontal_splitter.addWidget(log_panel)
        horizontal_splitter.setSizes([500, 500])  # Two panels equal size
        
        # Add splitter to content layout
        content_layout.addWidget(horizontal_splitter)
        
        # Add to main layout
        # main_tab değişkeni satır 520'de oluşturulmuş ve layout'u satır 521'de ayarlanmış
        main_tab_layout.addWidget(content_container, 1)  # 1 stretch factor
        
        # Add current main_tab to tabs widget
        self.tabs.addTab(main_tab, "Renamer - Organizer")
        
        # Search Term tab
        search_term_tab = QWidget()
        search_term_layout = QVBoxLayout(search_term_tab)
        search_term_layout.setContentsMargins(15, 15, 15, 15)
        
        # PDF Directory selection
        search_dir_group = QGroupBox("PDF Directory")
        search_dir_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(search_dir_group)
        search_dir_layout = QVBoxLayout(search_dir_group)
        search_dir_layout.setContentsMargins(6, 12, 6, 6)
        search_dir_layout.setSpacing(3)

        search_dir_input_layout = QHBoxLayout()
        self.search_dir_edit = QLineEdit()
        self.search_dir_edit.setPlaceholderText("Select the folder containing PDF files to search")
        search_dir_input_layout.addWidget(self.search_dir_edit)

        self.search_browse_btn = QPushButton("Browse...")
        self.search_browse_btn.clicked.connect(self.browse_search_directory)
        self.search_browse_btn.setStyleSheet(browse_button_style)
        search_dir_input_layout.addWidget(self.search_browse_btn)

        self.search_open_dir_btn = QPushButton("📁")
        self.search_open_dir_btn.setToolTip("Open the selected PDF directory")
        self.search_open_dir_btn.clicked.connect(lambda: self.open_folder(self.search_dir_edit.text()))
        self.search_open_dir_btn.setStyleSheet(browse_button_style)
        search_dir_input_layout.addWidget(self.search_open_dir_btn)

        search_dir_layout.addLayout(search_dir_input_layout)
        search_term_layout.addWidget(search_dir_group)

        # Search options group
        search_options_group = QGroupBox("Search Options")
        search_options_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(search_options_group)
        search_options_layout = QVBoxLayout(search_options_group)
        search_options_layout.setContentsMargins(6, 12, 6, 6)
        search_options_layout.setSpacing(10)
        
        # Main horizontal layout for all search options in a single row
        search_row_layout = QHBoxLayout()
        search_row_layout.setSpacing(15)
        
        # Add help button at the beginning
        help_button = QPushButton()
        help_button.setIcon(QIcon("resources/question.png"))
        help_button.setIconSize(QSize(24, 24))
        help_button.setToolTip("Show help for search options")
        help_button.setFixedSize(30, 30)
        help_button.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                border-radius: 15px;
            }
            QPushButton:hover {
                background-color: #E0E0E0;
            }
        """)
        help_button.clicked.connect(self.show_search_options_help)
        search_row_layout.addWidget(help_button)
        
        # Keyword input with increased width
        keyword_layout = QHBoxLayout()
        keyword_layout.setSpacing(5)
        keyword_label = QLabel("Keyword:")
        keyword_layout.addWidget(keyword_label)
        self.search_keyword_edit = QLineEdit()
        self.search_keyword_edit.setPlaceholderText("Enter search term")
        self.search_keyword_edit.setMinimumWidth(400)  # Double the previous width (was 200)
        keyword_layout.addWidget(self.search_keyword_edit)
        search_row_layout.addLayout(keyword_layout)
        
        # Divider line
        v_line = QFrame()
        v_line.setFrameShape(QFrame.VLine)
        v_line.setFrameShadow(QFrame.Sunken)
        v_line.setStyleSheet("background-color: #E0E0E0;")
        search_row_layout.addWidget(v_line)
        
        # Checkboxes reused icons
        uncheck_icon = "resources/unchecked.png"
        check_icon = "resources/check.png"
        
        # Options group label
        options_label = QLabel("Options:")
        search_row_layout.addWidget(options_label)
        
        # Exact match option - inline
        self.exact_match_check = QCheckBox("Exact Match")
        self.exact_match_check.setIcon(QIcon(uncheck_icon))
        self.exact_match_check.setChecked(False)
        self.exact_match_check.stateChanged.connect(
            lambda state: self.exact_match_check.setIcon(QIcon(check_icon if state else uncheck_icon))
        )
        search_row_layout.addWidget(self.exact_match_check)
        
        # Case sensitive option - inline
        self.case_match_check = QCheckBox("Case Sensitive")
        self.case_match_check.setIcon(QIcon(uncheck_icon))
        self.case_match_check.setChecked(False)
        self.case_match_check.stateChanged.connect(
            lambda state: self.case_match_check.setIcon(QIcon(check_icon if state else uncheck_icon))
        )
        search_row_layout.addWidget(self.case_match_check)
        
        # Regex option - inline
        self.regex_match_check = QCheckBox("Use Regex")
        self.regex_match_check.setIcon(QIcon(uncheck_icon))
        self.regex_match_check.setChecked(False)
        self.regex_match_check.stateChanged.connect(
            lambda state: self.regex_match_check.setIcon(QIcon(check_icon if state else uncheck_icon))
        )
        search_row_layout.addWidget(self.regex_match_check)
        
        # Add stretch to push everything to the left
        search_row_layout.addStretch(1)
        
        # Add the row layout to the main search options layout
        search_options_layout.addLayout(search_row_layout)
        
        # Add search options to tab
        search_term_layout.addWidget(search_options_group)

        # Buttons layout
        search_buttons_layout = QHBoxLayout()
        search_buttons_layout.setSpacing(20)  # Increase spacing between buttons
        
        self.search_start_btn = QPushButton("Start Search")
        self.search_start_btn.clicked.connect(self.start_search_processing)
        self.search_start_btn.setMinimumSize(350, 50)  # Larger and wider button
        self.search_start_btn.setStyleSheet("""
            QPushButton {
                background-color: #107C10;
                color: white;
                border: none;
                padding: 14px 30px;
                border-radius: 10px;
                font-weight: 600;
                font-size: 14px;
                text-align: center;
            }
            QPushButton:hover {
                background-color: #0A6A0A;
            }
            QPushButton:pressed {
                background-color: #095909;
            }
        """)
        search_buttons_layout.addWidget(self.search_start_btn, 3)  # 3:1 ratio for width
        
        self.search_stop_btn = QPushButton("Stop")
        self.search_stop_btn.clicked.connect(self.stop_search_processing)
        self.search_stop_btn.setEnabled(False)
        self.search_stop_btn.setMinimumSize(120, 50)
        self.search_stop_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {error_color};
                color: white;
                border: none;
                padding: 14px 20px;
                border-radius: 10px;
                font-weight: 600;
                font-size: 14px;
            }}
            QPushButton:hover {{
                background-color: #D31020;
            }}
            QPushButton:pressed {{
                background-color: #B80E1B;
            }}
            QPushButton:disabled {{
                background-color: {divider_color};
                color: {secondary_text};
            }}
        """)
        search_buttons_layout.addWidget(self.search_stop_btn, 1)  # 1 unit width
        
        search_term_layout.addLayout(search_buttons_layout)

        # Progress bar is removed
        # self.search_progress_bar = QProgressBar()
        # self.search_progress_bar.setTextVisible(True)
        # self.search_progress_bar.setValue(0)
        # self.search_progress_bar.setFormat("%p% (%v/%m files)")
        # search_term_layout.addWidget(self.search_progress_bar)

        # Results display
        search_results_group = QGroupBox("Search Results")
        search_results_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(search_results_group)
        search_results_layout = QVBoxLayout(search_results_group)
        search_results_layout.setContentsMargins(6, 12, 6, 6)

        # Results table
        self.search_results_table = QTableWidget()
        self.search_results_table.setColumnCount(7)  # Increased column count to add page number
        self.search_results_table.setHorizontalHeaderLabels(["DOI", "PDF Name", "Page", "Keyword", "Previous", "Matched", "Next"])
        self.search_results_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.search_results_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.search_results_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.search_results_table.setAlternatingRowColors(True)
        search_results_layout.addWidget(self.search_results_table)

        search_results_buttons = QHBoxLayout()
        self.save_search_results_btn = QPushButton("Save Results")
        self.save_search_results_btn.clicked.connect(self.save_search_results)
        self.save_search_results_btn.setStyleSheet(browse_button_style)
        self.save_search_results_btn.setEnabled(False)
        search_results_buttons.addWidget(self.save_search_results_btn)

        self.clear_search_results_btn = QPushButton("Clear Results")
        self.clear_search_results_btn.clicked.connect(self.clear_search_results)
        self.clear_search_results_btn.setStyleSheet(browse_button_style)
        search_results_buttons.addWidget(self.clear_search_results_btn)
        search_results_layout.addLayout(search_results_buttons)

        search_term_layout.addWidget(search_results_group)

        # Add the tab
        self.tabs.addTab(search_term_tab, "Search Keywords")
        
        # Create API settings tab
        api_tab = QWidget()
        api_layout = QVBoxLayout(api_tab)
        api_layout.setContentsMargins(15, 15, 15, 15)  # Wider margins
        api_layout.setSpacing(15)  # More vertical spacing

        # Custom style for API settings
        api_tab.setStyleSheet("""
            QLabel { 
                font-size: 9.5pt; 
            }
            QGroupBox { 
                font-size: 10.5pt;
            }
        """)
        
        # API checkbox style
        api_checkbox_style = f"""
            QCheckBox {{
                padding: 8px;
                font-size: 11pt;
                color: {text_color};
            }}
            QCheckBox::indicator {{
                width: 22px;
                height: 22px;
                border-radius: 4px;
            }}
            QCheckBox::indicator:unchecked {{
                image: url(resources/unchecked.png);
            }}
            QCheckBox::indicator:checked {{
                image: url(resources/check.png);
            }}
            QCheckBox::indicator:unchecked:hover {{
                border: 1px solid {primary_color};
            }}
        """
        
        # Free APIs (no configuration required)
        free_apis_group = QGroupBox("Free APIs (Always Enabled)")
        free_apis_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(free_apis_group)
        free_apis_layout = QVBoxLayout(free_apis_group)
        
        # Free APIs explanation
        free_explanation = QLabel(
            "<i>These APIs don't require any registration or authentication. "
            "They are always available and provide metadata for your documents.</i>"
        )
        free_explanation.setWordWrap(True)
        free_apis_layout.addWidget(free_explanation)
        
        # Free APIs grid for checkboxes
        free_apis_grid = QGridLayout()
        
        # CrossRef
        self.crossref_check = QCheckBox("CrossRef")
        self.crossref_check.setChecked(True)
        self.crossref_check.setEnabled(False)  # Always active
        self.crossref_check.setToolTip("CrossRef API is always enabled (free, no registration required)")
        self.crossref_check.setStyleSheet(api_checkbox_style)
        free_apis_grid.addWidget(self.crossref_check, 0, 0)
        
        # DataCite
        self.datacite_check = QCheckBox("DataCite")
        self.datacite_check.setChecked(True)
        self.datacite_check.setEnabled(False)  # Always active
        self.datacite_check.setToolTip("DataCite API is always enabled (free, no registration required)")
        self.datacite_check.setStyleSheet(api_checkbox_style)
        free_apis_grid.addWidget(self.datacite_check, 0, 1)
        
        # Europe PMC
        self.europepmc_check = QCheckBox("Europe PMC")
        self.europepmc_check.setChecked(True)
        self.europepmc_check.setEnabled(False)  # Always active
        self.europepmc_check.setToolTip("Europe PMC API is always enabled (free, no registration required)")
        self.europepmc_check.setStyleSheet(api_checkbox_style)
        free_apis_grid.addWidget(self.europepmc_check, 1, 0)
        
        # Semantic Scholar
        self.semantic_scholar_check = QCheckBox("Semantic Scholar")
        self.semantic_scholar_check.setChecked(True)
        self.semantic_scholar_check.setEnabled(False)  # Always active
        self.semantic_scholar_check.setToolTip("Semantic Scholar API is always enabled (free, no registration required)")
        self.semantic_scholar_check.setStyleSheet(api_checkbox_style)
        free_apis_grid.addWidget(self.semantic_scholar_check, 1, 1)
        
        # OpenAlex
        self.openalex_check = QCheckBox("OpenAlex")
        self.openalex_check.setChecked(True)
        self.openalex_check.setEnabled(False)  # Always active
        self.openalex_check.setToolTip("OpenAlex API is always enabled (free, no registration required)")
        self.openalex_check.setStyleSheet(api_checkbox_style)
        free_apis_grid.addWidget(self.openalex_check, 2, 0)
        
        # Unpaywall
        self.unpaywall_check = QCheckBox("Unpaywall")
        self.unpaywall_check.setChecked(True)
        self.unpaywall_check.setEnabled(False)  # Always active
        self.unpaywall_check.setToolTip("Unpaywall API is always enabled (free, no registration required)")
        self.unpaywall_check.setStyleSheet(api_checkbox_style)
        free_apis_grid.addWidget(self.unpaywall_check, 2, 1)
        
        # Add Free API grid to free_apis_layout
        free_apis_layout.addLayout(free_apis_grid)
        
        # Optional Email configuration
        email_apis_group = QGroupBox("Optional Email Configuration")
        email_apis_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(email_apis_group)
        email_apis_layout = QVBoxLayout(email_apis_group)
        
        # Email usage explanation
        email_explanation = QLabel(
            "<i>Some APIs work better with an email address. Providing your email can improve rate limits "
            "and help comply with API terms of service.</i>"
        )
        email_explanation.setWordWrap(True)
        email_apis_layout.addWidget(email_explanation)
        
        # OpenAlex Email
        openalex_layout = QHBoxLayout()
        openalex_email_label = QLabel("OpenAlex Email (optional):")
        openalex_layout.addWidget(openalex_email_label)
        
        self.openalex_email_edit = QLineEdit()
        self.openalex_email_edit.setPlaceholderText("Your email address for better rate limits")
        openalex_layout.addWidget(self.openalex_email_edit)
        
        email_apis_layout.addLayout(openalex_layout)
        
        # Unpaywall Email
        unpaywall_layout = QHBoxLayout()
        unpaywall_email_label = QLabel("Unpaywall Email (optional):")
        unpaywall_layout.addWidget(unpaywall_email_label)
        
        self.unpaywall_email_edit = QLineEdit()
        self.unpaywall_email_edit.setPlaceholderText("Your email address for better rate limits")
        unpaywall_layout.addWidget(self.unpaywall_email_edit)
        
        email_apis_layout.addLayout(unpaywall_layout)
        
        # Semantic Scholar API Key (optional)
        semantic_scholar_layout = QHBoxLayout()
        semantic_scholar_key_label = QLabel("Semantic Scholar API Key (optional):")
        semantic_scholar_layout.addWidget(semantic_scholar_key_label)
        
        self.semantic_scholar_key_edit = QLineEdit()
        self.semantic_scholar_key_edit.setPlaceholderText("For higher rate limits if you have a key")
        self.semantic_scholar_key_edit.setEchoMode(QLineEdit.Password)
        semantic_scholar_layout.addWidget(self.semantic_scholar_key_edit)
        
        email_apis_layout.addLayout(semantic_scholar_layout)
        
        # API Keys group - Only for Scopus now
        api_keys_group = QGroupBox("APIs Requiring API Keys")
        api_keys_group.setStyleSheet(self.groupbox_style)
        self.add_shadow_effect(api_keys_group)
        api_keys_layout = QVBoxLayout(api_keys_group)
        
        # API keys explanation
        apikey_explanation = QLabel(
            "<i>These APIs <b>require registration</b> to obtain an API key. "
            "The Scopus API will not work without a valid API key.</i>"
        )
        apikey_explanation.setWordWrap(True)
        api_keys_layout.addWidget(apikey_explanation)
        
        # Scopus
        scopus_layout = QHBoxLayout()
        
        # Scopus checkbox
        self.scopus_check = QCheckBox("Scopus")
        self.scopus_check.setToolTip("Enable Scopus API for metadata retrieval (requires registration)")
        self.scopus_check.setStyleSheet(api_checkbox_style)
        scopus_layout.addWidget(self.scopus_check)
        
        scopus_key_label = QLabel("API Key:")
        scopus_layout.addWidget(scopus_key_label)
        
        self.scopus_key_edit = QLineEdit()
        self.scopus_key_edit.setPlaceholderText("Register at dev.elsevier.com to get an API key")
        self.scopus_key_edit.setEchoMode(QLineEdit.Password)
        scopus_layout.addWidget(self.scopus_key_edit)
        
        api_keys_layout.addLayout(scopus_layout)
        
        # Save button
        save_api_btn = QPushButton("Save API Settings")
        save_api_btn.clicked.connect(self.save_api_settings)
        save_api_btn.setStyleSheet("""
            QPushButton {
                background-color: #107C10;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 8px;
                font-weight: 600;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #0A6A0A;
            }
            QPushButton:pressed {
                background-color: #095909;
            }
        """)
        
        # Make API tab larger (content expanded)
        free_apis_group.setMinimumHeight(180)  # Increase default height
        email_apis_group.setMinimumHeight(140)
        api_keys_group.setMinimumHeight(140)
        
        # Add API groups to layout
        api_layout.addWidget(free_apis_group)
        api_layout.addWidget(email_apis_group)
        api_layout.addWidget(api_keys_group)
        api_layout.addWidget(save_api_btn, alignment=Qt.AlignCenter)
        api_layout.addStretch(1)  # Add space
        
        # Add the API tab to the tab widget
        self.tabs.addTab(api_tab, "API Settings")
        
        # Add Statistics Tab (General)
        statistics_tab = QWidget()
        statistics_layout = QVBoxLayout(statistics_tab)
        statistics_layout.setContentsMargins(15, 15, 15, 15)
        statistics_layout.setSpacing(15)
        
        # Statistics tab welcome message and description
        stats_header = QLabel("Process Statistics")
        stats_header.setFont(QFont("Segoe UI", 14, QFont.Bold))
        stats_header.setStyleSheet(f"color: {primary_color};")
        stats_header.setAlignment(Qt.AlignCenter)
        statistics_layout.addWidget(stats_header)
        
        stats_description = QLabel(
            "This tab displays statistics about your PDF processing results. "
            "Statistics are automatically updated when processing is completed."
        )
        stats_description.setWordWrap(True)
        stats_description.setStyleSheet(f"color: {secondary_text}; font-size: 10pt;")
        stats_description.setAlignment(Qt.AlignCenter)
        statistics_layout.addWidget(stats_description)
        
        # Dashboard layout for statistics
        stats_dashboard = QWidget()
        dashboard_layout = QGridLayout(stats_dashboard)
        dashboard_layout.setSpacing(20)
        
        # Create statistic cards
        # Card 1 - General Stats
        general_stats_card = QGroupBox("General Statistics")
        general_stats_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(general_stats_card)
        general_stats_layout = QVBoxLayout(general_stats_card)
        general_stats_layout.setContentsMargins(10, 15, 10, 15)
        
        # General stats content
        self.stat_total_files = QLabel("Total Files Processed: 0")
        self.stat_total_files.setFont(QFont("Segoe UI", 10))
        general_stats_layout.addWidget(self.stat_total_files)
        
        self.stat_renamed_files = QLabel("Successfully Renamed: 0")
        self.stat_renamed_files.setFont(QFont("Segoe UI", 10))
        self.stat_renamed_files.setStyleSheet(f"color: {success_color};")
        general_stats_layout.addWidget(self.stat_renamed_files)
        
        self.stat_problematic_files = QLabel("Files Not Renamed: 0")
        self.stat_problematic_files.setFont(QFont("Segoe UI", 10))
        self.stat_problematic_files.setStyleSheet(f"color: {error_color};")
        general_stats_layout.addWidget(self.stat_problematic_files)
        
        self.stat_success_rate = QLabel("Success Rate: 0%")
        self.stat_success_rate.setFont(QFont("Segoe UI", 10, QFont.Bold))
        general_stats_layout.addWidget(self.stat_success_rate)
        
        dashboard_layout.addWidget(general_stats_card, 0, 0)
        
        # Card 2 - Performance Stats
        performance_stats_card = QGroupBox("Performance Metrics")
        performance_stats_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(performance_stats_card)
        performance_stats_layout = QVBoxLayout(performance_stats_card)
        performance_stats_layout.setContentsMargins(10, 15, 10, 15)
        
        # Performance stats content
        self.stat_total_time = QLabel("Total Processing Time: 0 seconds")
        self.stat_total_time.setFont(QFont("Segoe UI", 10))
        performance_stats_layout.addWidget(self.stat_total_time)
        
        # Add progress bar for total time visualization
        self.total_time_progress = QProgressBar()
        self.total_time_progress.setValue(0)
        self.total_time_progress.setTextVisible(False)
        self.total_time_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #0078D4;
                border-radius: 4px;
            }}
        """)
        performance_stats_layout.addWidget(self.total_time_progress)
        
        self.stat_per_file_time = QLabel("Average Time Per File: 0 seconds")
        self.stat_per_file_time.setFont(QFont("Segoe UI", 10))
        performance_stats_layout.addWidget(self.stat_per_file_time)
        
        # Add progress bar for per file time visualization
        self.per_file_time_progress = QProgressBar()
        self.per_file_time_progress.setValue(0)
        self.per_file_time_progress.setTextVisible(False)
        self.per_file_time_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #50ABF1;
                border-radius: 4px;
            }}
        """)
        performance_stats_layout.addWidget(self.per_file_time_progress)
        
        self.stat_processing_speed = QLabel("Processing Speed: 0 files/second")
        self.stat_processing_speed.setFont(QFont("Segoe UI", 10))
        performance_stats_layout.addWidget(self.stat_processing_speed)
        
        # Add progress bar for processing speed visualization
        self.processing_speed_progress = QProgressBar()
        self.processing_speed_progress.setValue(0)
        self.processing_speed_progress.setTextVisible(False)
        self.processing_speed_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #107C10;
                border-radius: 4px;
            }}
        """)
        performance_stats_layout.addWidget(self.processing_speed_progress)
        
        # Memory usage estimate (simplified)
        self.stat_memory_usage = QLabel("Estimated Memory Usage: 0 MB")
        self.stat_memory_usage.setFont(QFont("Segoe UI", 10))
        performance_stats_layout.addWidget(self.stat_memory_usage)
        
        # Add progress bar for memory usage visualization
        self.memory_usage_progress = QProgressBar()
        self.memory_usage_progress.setValue(0)
        self.memory_usage_progress.setTextVisible(False)
        self.memory_usage_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #F8A800;
                border-radius: 4px;
            }}
        """)
        performance_stats_layout.addWidget(self.memory_usage_progress)
        
        dashboard_layout.addWidget(performance_stats_card, 0, 1)
        
        # Card 3 - Accuracy Metrics
        accuracy_metrics_card = QGroupBox("Accuracy Metrics")
        accuracy_metrics_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(accuracy_metrics_card)
        accuracy_metrics_layout = QVBoxLayout(accuracy_metrics_card)
        accuracy_metrics_layout.setContentsMargins(10, 15, 10, 15)
        
        # Accuracy metrics content
        self.stat_doi_detection = QLabel("DOI Detection Rate: 0%")
        self.stat_doi_detection.setFont(QFont("Segoe UI", 10))
        accuracy_metrics_layout.addWidget(self.stat_doi_detection)
        
        # Add DOI detection rate progress bar
        self.doi_detection_progress = QProgressBar()
        self.doi_detection_progress.setValue(0)
        self.doi_detection_progress.setTextVisible(False)
        self.doi_detection_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #0078D4;
                border-radius: 4px;
            }}
        """)
        accuracy_metrics_layout.addWidget(self.doi_detection_progress)
        
        self.stat_metadata_quality = QLabel("Metadata Quality: 0%")
        self.stat_metadata_quality.setFont(QFont("Segoe UI", 10))
        accuracy_metrics_layout.addWidget(self.stat_metadata_quality)
        
        # Add metadata quality progress bar
        self.metadata_quality_progress = QProgressBar()
        self.metadata_quality_progress.setValue(0)
        self.metadata_quality_progress.setTextVisible(False)
        self.metadata_quality_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #107C10;
                border-radius: 4px;
            }}
        """)
        accuracy_metrics_layout.addWidget(self.metadata_quality_progress)
        
        self.stat_categorization_quality = QLabel("Categorization Quality: 0%")
        self.stat_categorization_quality.setFont(QFont("Segoe UI", 10))
        accuracy_metrics_layout.addWidget(self.stat_categorization_quality)
        
        # Add categorization quality progress bar
        self.categorization_quality_progress = QProgressBar()
        self.categorization_quality_progress.setValue(0)
        self.categorization_quality_progress.setTextVisible(False)
        self.categorization_quality_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #50ABF1;
                border-radius: 4px;
            }}
        """)
        accuracy_metrics_layout.addWidget(self.categorization_quality_progress)
        
        # Time savings estimate
        self.stat_time_savings = QLabel("Estimated Time Saved: 0 minutes")
        self.stat_time_savings.setFont(QFont("Segoe UI", 10))
        accuracy_metrics_layout.addWidget(self.stat_time_savings)
        
        # Add time savings progress bar
        self.time_savings_progress = QProgressBar()
        self.time_savings_progress.setValue(0)
        self.time_savings_progress.setTextVisible(False)
        self.time_savings_progress.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: #F0F0F0;
                border-radius: 4px;
                height: 8px;
                margin-bottom: 12px;
            }}
            QProgressBar::chunk {{
                background-color: #F8A800;
                border-radius: 4px;
            }}
        """)
        accuracy_metrics_layout.addWidget(self.time_savings_progress)
        
        dashboard_layout.addWidget(accuracy_metrics_card, 1, 0)
        
        # Card 4 - API Statistics
        api_stats_card = QGroupBox("API Usage Statistics")
        api_stats_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(api_stats_card)
        api_stats_layout = QVBoxLayout(api_stats_card)
        api_stats_layout.setContentsMargins(10, 15, 10, 15)
        
        # API stats list
        self.api_stats_list = QListWidget()
        self.api_stats_list.setStyleSheet(f"""
            QListWidget {{
                border: none;
                background-color: {card_color};
            }}
            QListWidget::item {{
                padding: 5px;
                border-radius: 5px;
                background-color: {card_color};
            }}
            QListWidget::item:hover {{
                background-color: {hover_color};
            }}
            QListWidget::item:selected {{
                background-color: #E3F2FD;
                color: {text_color};
            }}
        """)
        api_stats_layout.addWidget(self.api_stats_list)
        
        dashboard_layout.addWidget(api_stats_card, 1, 1)
        
        # Add main dashboard to layout
        statistics_layout.addWidget(stats_dashboard)
        
        # Error breakdown section
        error_section = QGroupBox("Error Analysis")
        error_section.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(error_section)
        error_layout = QVBoxLayout(error_section)
        
        # Error breakdown list
        self.error_breakdown_list = QListWidget()
        self.error_breakdown_list.setStyleSheet(f"""
            QListWidget {{
                border: none;
                background-color: {card_color};
            }}
            QListWidget::item {{
                padding: 5px;
                border-radius: 5px;
                background-color: {card_color};
            }}
            QListWidget::item:hover {{
                background-color: {hover_color};
            }}
            QListWidget::item:selected {{
                background-color: #E3F2FD;
                color: {text_color};
            }}
        """)
        self.error_breakdown_list.addItem("No errors identified yet")
        error_layout.addWidget(self.error_breakdown_list)
        
        statistics_layout.addWidget(error_section)
        
        # Add a placeholder message for when no data is available
        self.stats_placeholder = QLabel("No statistics data available yet. Complete the processing to view statistics.")
        self.stats_placeholder.setAlignment(Qt.AlignCenter)
        self.stats_placeholder.setStyleSheet(f"color: {secondary_text}; font-size: 11pt; margin: 20px;")
        self.stats_placeholder.setWordWrap(True)
        statistics_layout.addWidget(self.stats_placeholder)
        
        # Add the statistics tab to the tab widget
        self.tabs.addTab(statistics_tab, "General Statistics")
        
        # Add Publication Statistics Tab
        pub_stats_tab = QWidget()
        pub_stats_layout = QVBoxLayout(pub_stats_tab)
        pub_stats_layout.setContentsMargins(15, 15, 15, 15)
        pub_stats_layout.setSpacing(15)
        
        # Tab header
        pub_stats_header = QLabel("Publication Statistics")
        pub_stats_header.setFont(QFont("Segoe UI", 14, QFont.Bold))
        pub_stats_header.setStyleSheet(f"color: {primary_color};")
        pub_stats_header.setAlignment(Qt.AlignCenter)
        pub_stats_layout.addWidget(pub_stats_header)
        
        # Tab description
        pub_stats_description = QLabel(
            "This tab shows detailed statistics about publications, authors, journals, and years."
        )
        pub_stats_description.setWordWrap(True)
        pub_stats_description.setStyleSheet(f"color: {secondary_text}; font-size: 10pt;")
        pub_stats_description.setAlignment(Qt.AlignCenter)
        pub_stats_layout.addWidget(pub_stats_description)
        
        # Publication stats grid layout
        pub_stats_grid = QWidget()
        pub_grid_layout = QGridLayout(pub_stats_grid)
        pub_grid_layout.setSpacing(20)
        
        # Categorization Stats Card
        categorization_stats_card = QGroupBox("Categorization Statistics")
        categorization_stats_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(categorization_stats_card)
        categorization_stats_layout = QVBoxLayout(categorization_stats_card)
        categorization_stats_layout.setContentsMargins(10, 15, 10, 15)
        
        # Categorization stats content
        self.stat_categorized_by_journal = QLabel("Categorized by Journal: 0 files")
        self.stat_categorized_by_journal.setFont(QFont("Segoe UI", 10))
        categorization_stats_layout.addWidget(self.stat_categorized_by_journal)
        
        self.stat_categorized_by_author = QLabel("Categorized by Author: 0 files")
        self.stat_categorized_by_author.setFont(QFont("Segoe UI", 10))
        categorization_stats_layout.addWidget(self.stat_categorized_by_author)
        
        self.stat_categorized_by_year = QLabel("Categorized by Year: 0 files")
        self.stat_categorized_by_year.setFont(QFont("Segoe UI", 10))
        categorization_stats_layout.addWidget(self.stat_categorized_by_year)
        
        self.stat_categorized_by_subject = QLabel("Categorized by Subject: 0 files")
        self.stat_categorized_by_subject.setFont(QFont("Segoe UI", 10))
        categorization_stats_layout.addWidget(self.stat_categorized_by_subject)
        
        pub_grid_layout.addWidget(categorization_stats_card, 0, 0)
        
        # Top Authors Card
        top_authors_card = QGroupBox("Authors")
        top_authors_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(top_authors_card)
        top_authors_layout = QVBoxLayout(top_authors_card)
        top_authors_layout.setContentsMargins(10, 15, 10, 15)
        
        # Authors list
        self.top_authors_list = QListWidget()
        self.top_authors_list.setStyleSheet(f"""
            QListWidget {{
                border: none;
                background-color: {card_color};
            }}
            QListWidget::item {{
                padding: 8px;
                border-radius: 5px;
                background-color: {card_color};
                margin-bottom: 2px;
            }}
            QListWidget::item:hover {{
                background-color: {hover_color};
            }}
            QListWidget::item:selected {{
                background-color: #E3F2FD;
                color: {text_color};
            }}
        """)
        self.top_authors_list.addItem("No data available yet")
        top_authors_layout.addWidget(self.top_authors_list)
        
        pub_grid_layout.addWidget(top_authors_card, 0, 1)
        
        # Top Journals Card
        top_journals_card = QGroupBox("Journals")
        top_journals_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(top_journals_card)
        top_journals_layout = QVBoxLayout(top_journals_card)
        top_journals_layout.setContentsMargins(10, 15, 10, 15)
        
        # Journals list
        self.top_journals_list = QListWidget()
        self.top_journals_list.setStyleSheet(f"""
            QListWidget {{
                border: none;
                background-color: {card_color};
            }}
            QListWidget::item {{
                padding: 8px;
                border-radius: 5px;
                background-color: {card_color};
                margin-bottom: 2px;
            }}
            QListWidget::item:hover {{
                background-color: {hover_color};
            }}
            QListWidget::item:selected {{
                background-color: #E3F2FD;
                color: {text_color};
            }}
        """)
        self.top_journals_list.addItem("No data available yet")
        top_journals_layout.addWidget(self.top_journals_list)
        
        pub_grid_layout.addWidget(top_journals_card, 1, 0)
        
        # Years Distribution Card
        years_distribution_card = QGroupBox("Year Distribution")
        years_distribution_card.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(years_distribution_card)
        years_distribution_layout = QVBoxLayout(years_distribution_card)
        years_distribution_layout.setContentsMargins(10, 15, 10, 15)
        
        # Years list
        self.year_distribution_list = QListWidget()
        self.year_distribution_list.setStyleSheet(f"""
            QListWidget {{
                border: none;
                background-color: {card_color};
            }}
            QListWidget::item {{
                padding: 8px;
                border-radius: 5px;
                background-color: {card_color};
                margin-bottom: 2px;
            }}
            QListWidget::item:hover {{
                background-color: {hover_color};
            }}
            QListWidget::item:selected {{
                background-color: #E3F2FD;
                color: {text_color};
            }}
        """)
        self.year_distribution_list.addItem("No data available yet")
        years_distribution_layout.addWidget(self.year_distribution_list)
        
        pub_grid_layout.addWidget(years_distribution_card, 1, 1)
        
        # Add main grid to layout
        pub_stats_layout.addWidget(pub_stats_grid)
        
        # Subject Distribution Section
        subject_section = QGroupBox("Subject Distribution")
        subject_section.setStyleSheet(f"""
            {self.groupbox_style}
            QGroupBox {{
                background-color: {card_color};
            }}
        """)
        self.add_shadow_effect(subject_section)
        subject_layout = QVBoxLayout(subject_section)
        
        # Subject list
        self.subject_list = QListWidget()
        self.subject_list.setStyleSheet(f"""
            QListWidget {{
                border: none;
                background-color: {card_color};
            }}
            QListWidget::item {{
                padding: 8px;
                border-radius: 5px;
                background-color: {card_color};
            }}
            QListWidget::item:hover {{
                background-color: {hover_color};
            }}
            QListWidget::item:selected {{
                background-color: #E3F2FD;
                color: {text_color};
            }}
        """)
        self.subject_list.addItem("No subject data available yet")
        subject_layout.addWidget(self.subject_list)
        
        pub_stats_layout.addWidget(subject_section)
        
        # Add publication statistics tab
        self.tabs.addTab(pub_stats_tab, "Publication Statistics")
        
        # Set the main tab as the default
        self.tabs.setCurrentIndex(0)
        
        # Simply modify the text to include extra spaces for tabs to make them wider
        self.tabs.setTabText(0, "   Renamer - Organizer   ")  # Main tab with more space
        search_tab_index = 1  # Search Keywords tab index (0-based)
        self.tabs.setTabText(search_tab_index, "   Search Keywords   ")
        self.tabs.setTabText(2, "   API Settings   ")  # API settings tab with more space
        self.tabs.setTabText(3, "   General Statistics   ")  # General statistics tab
        self.tabs.setTabText(4, "   Publication Statistics   ")  # Publication statistics tab
        
        # Windows 11 style shadows for panels
        self.add_shadow_effect(results_panel)
        self.add_shadow_effect(log_panel)
    
    def clear_status_list(self):
        """
        Clear the status list.
        """
        self.results_list.clear()
    
    def load_settings(self):
        """
        Load application settings.
        """
        settings = QSettings("LitOrganizer", "LitOrganizer")
        
        # Main settings
        self.dir_edit.setText(settings.value("pdf_directory", ""))
        self.unclassified_dir_edit.setText(settings.value("unclassified_directory", ""))
        self.use_ocr_check.setChecked(settings.value("use_ocr", False, type=bool))
        self.create_backups_check.setChecked(settings.value("create_backups", True, type=bool))
        self.move_problematic_check.setChecked(settings.value("move_problematic", True, type=bool))
        self.create_references_check.setChecked(settings.value("create_references", True, type=bool))
        
        # Categorization settings
        self.categorize_by_journal_check.setChecked(settings.value("categorize_by_journal", False, type=bool))
        self.categorize_by_author_check.setChecked(settings.value("categorize_by_author", False, type=bool))
        self.categorize_by_year_check.setChecked(settings.value("categorize_by_year", False, type=bool))
        self.categorize_by_subject_check.setChecked(settings.value("categorize_by_subject", False, type=bool))
        
        # API settings
        self.load_api_settings()
    
    def save_settings(self):
        """
        Save application settings.
        """
        settings = QSettings("LitOrganizer", "LitOrganizer")
        
        # Main settings
        settings.setValue("pdf_directory", self.dir_edit.text())
        settings.setValue("unclassified_directory", self.unclassified_dir_edit.text())
        settings.setValue("use_ocr", self.use_ocr_check.isChecked())
        settings.setValue("create_backups", self.create_backups_check.isChecked())
        settings.setValue("move_problematic", self.move_problematic_check.isChecked())
        settings.setValue("create_references", self.create_references_check.isChecked())
        
        # Categorization settings
        settings.setValue("categorize_by_journal", self.categorize_by_journal_check.isChecked())
        settings.setValue("categorize_by_author", self.categorize_by_author_check.isChecked())
        settings.setValue("categorize_by_year", self.categorize_by_year_check.isChecked())
        settings.setValue("categorize_by_subject", self.categorize_by_subject_check.isChecked())

    def load_api_settings(self):
        """
        Load API settings from the config file.
        """
        try:
            config = load_api_config()
            
            # Free APIs - always enabled
            self.crossref_check.setChecked(True)
            self.openalex_check.setChecked(True)
            self.datacite_check.setChecked(True)
            self.europepmc_check.setChecked(True)
            self.semantic_scholar_check.setChecked(True)
            self.unpaywall_check.setChecked(True)
            
            # Optional email settings
            self.openalex_email_edit.setText(config.get("openalex", {}).get("email", ""))
            self.unpaywall_email_edit.setText(config.get("unpaywall", {}).get("email", ""))
            self.semantic_scholar_key_edit.setText(config.get("semantic_scholar", {}).get("api_key", ""))
            
            # APIs requiring authentication
            self.scopus_check.setChecked(config.get("scopus", {}).get("enabled", False))
            self.scopus_key_edit.setText(config.get("scopus", {}).get("api_key", ""))
            
            self.logger.debug("Loaded API settings from config file")
        except Exception as e:
            self.logger.error(f"Error loading API settings: {e}")

    def save_api_settings(self):
        """
        Save API settings to the config file.
        """
        from pathlib import Path
        import json
        
        try:
            config_path = Path(__file__).parent.parent.parent / 'config' / 'api_keys.json'
            
            # Create directory if it doesn't exist
            config_path.parent.mkdir(exist_ok=True)
            
            config = {
                "crossref": {
                    "enabled": True  # Always enabled
                },
                "openalex": {
                    "email": self.openalex_email_edit.text(),
                    "enabled": True  # Always enabled
                },
                "datacite": {
                    "enabled": True  # Always enabled
                },
                "europepmc": {
                    "enabled": True  # Always enabled
                },
                "scopus": {
                    "api_key": self.scopus_key_edit.text(),
                    "enabled": self.scopus_check.isChecked()
                },
                "semantic_scholar": {
                    "api_key": self.semantic_scholar_key_edit.text(),
                    "enabled": True  # Always enabled
                },
                "unpaywall": {
                    "email": self.unpaywall_email_edit.text(),
                    "enabled": True  # Always enabled
                }
            }
            
            with open(config_path, 'w') as f:
                json.dump(config, f, indent=4)
            
            self.logger.info("API settings saved successfully")
            QMessageBox.information(self, "Settings Saved", "API settings have been saved successfully.")
        except Exception as e:
            self.logger.error(f"Error saving API settings: {e}")
            QMessageBox.critical(self, "Error", f"Failed to save API settings: {e}")

    def browse_directory(self):
        """
        Open a file dialog to select a directory.
        """
        directory = QFileDialog.getExistingDirectory(
            self,
            "Select Directory",
            self.dir_edit.text(),
            QFileDialog.ShowDirsOnly
        )
        
        if directory:
            self.dir_edit.setText(directory)
    
    def browse_unclassified_directory(self):
        """
        Open a file dialog to select an Unnamed Article files directory.
        """
        directory = QFileDialog.getExistingDirectory(
            self,
            "Select Unnamed Article Files Directory",
            self.unclassified_dir_edit.text() or self.dir_edit.text(),
            QFileDialog.ShowDirsOnly
        )
        
        if directory:
            self.unclassified_dir_edit.setText(directory)
    
    def start_processing(self):
        """
        Start processing PDF files.
        """
        # Check if a directory is selected
        directory = self.dir_edit.text()
        if not directory:
            QMessageBox.warning(self, "Error", "Please select a directory containing PDF files.")
            return
        
        # Create Path object
        directory_path = Path(directory)
        if not directory_path.exists() or not directory_path.is_dir():
            QMessageBox.warning(self, "Error", f"The directory {directory} does not exist.")
            return
        
        # Set unclassified directory
        unclassified_dir = self.unclassified_dir_edit.text()
        if not unclassified_dir:
            # Create an 'Unnamed Article' subfolder in the same directory
            unclassified_dir = str(directory_path / "Unnamed Article")
        self.unclassified_dir = unclassified_dir
        
        # Get categorization settings
        categorize_options = {
            "by_journal": self.categorize_by_journal_check.isChecked(),
            "by_author": self.categorize_by_author_check.isChecked(),
            "by_year": self.categorize_by_year_check.isChecked(),
            "by_subject": self.categorize_by_subject_check.isChecked()
        }
        
        # Check if there are PDF files in the directory
        pdf_files = list(directory_path.glob("*.pdf"))
        if not pdf_files:
            QMessageBox.warning(self, "Warning", f"No PDF files found in {directory}")
            return
        
        # Save settings
        self.save_settings()
        
        # Clear existing status
        self.clear_status_list()
        self.results_list.addItem(f"Found {len(pdf_files)} PDF files to process")
        
        # Record start time for performance tracking
        self.process_start_time = time.time()
        
        # Start worker thread
        self.worker_thread = WorkerThread(
            directory_path,
            use_ocr=self.use_ocr_check.isChecked(),
            create_references=self.create_references_check.isChecked(),
            create_backups=self.create_backups_check.isChecked(),
            move_problematic=self.move_problematic_check.isChecked(),
            problematic_dir=self.unclassified_dir,
            categorize_options=categorize_options,
            logger=self.logger
        )
        
        # Connect signals
        self.worker_thread.progress_update.connect(self.log_message)
        self.worker_thread.file_processed.connect(self.update_file_status)
        self.worker_thread.processing_complete.connect(self.processing_completed)
        self.worker_thread.error_occurred.connect(self.handle_error)
        self.worker_thread.progress_percentage.connect(self.update_progress_percentage)
        
        # Update UI - Start button appearance
        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.start_btn.setText("Processing... 0%")  # Progress indicator on Start button
        
        # Start thread
        self.worker_thread.start()
    
    def stop_processing(self):
        """
        Stop PDF processing.
        """
        if self.worker_thread and self.worker_thread.isRunning():
            self.log_message("Stopping processing... (This may take a moment)")
            
            # Progress bar pause
            self.start_btn.setText("Stopping...")
            self.start_btn.setStyleSheet("""
                QPushButton {
                    background-color: #F8A800;
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }
            """)
            
            self.worker_thread.terminate()
            
            # Disable the stop button to prevent multiple clicks
            self.stop_btn.setEnabled(False)
    
    def log_message(self, message):
        """
        Add a message to the log text area.
        
        Args:
            message (str): Message to add
        """
        self.log_text.append(message)
        # Scroll to bottom
        cursor = self.log_text.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.log_text.setTextCursor(cursor)
        
        # Process UI events
        QApplication.processEvents()
    
    def clear_log(self):
        """
        Clear the log text area.
        """
        self.log_text.clear()
    
    def update_file_status(self, filename, success):
        """
        Update the status of a file in the results list.
        
        Args:
            filename (str): Filename
            success (bool): Whether the file was processed successfully
        """
        item_text = f"{filename} - {'✅ Successfully Renamed' if success else '❌ Not Renamed (Problem)'}"
        item = QListWidgetItem(item_text)
        if success:
            item.setForeground(Qt.darkGreen)
        else:
            item.setForeground(Qt.darkRed)
        self.results_list.addItem(item)
        self.results_list.scrollToBottom()
    
    def processing_completed(self, processed, renamed, problematic):
        """
        Handle processing completion.
        
        Args:
            processed (int): Number of processed files
            renamed (int): Number of renamed files
            problematic (int): Number of unclassified files
        """
        # Record end time for performance tracking
        self.process_end_time = time.time()
        
        # Update UI
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        
        # Progress bar mark as completed
        self.update_progress_percentage(100)
        
        # Create completion message
        summary_item = QListWidgetItem("✅ Processing completed!")
        font = summary_item.font()
        font.setBold(True)
        summary_item.setFont(font)
        self.results_list.addItem(summary_item)
        
        stats_item = QListWidgetItem(f"Total Processed: {processed} | Successfully Renamed: {renamed} | Not Renamed: {problematic}")
        self.results_list.addItem(stats_item)
        self.results_list.scrollToBottom()
        
        # Update statistics tab
        self.update_statistics(processed, renamed, problematic)
        
        # Switch to statistics tab instead of showing message box
        self.tabs.setCurrentIndex(3)  # Index 3 is the Statistics tab
        
        # Log completion
        self.logger.info(f"Processing completed: {processed} files processed, {renamed} renamed, {problematic} problematic")
    
    def handle_error(self, error_message):
        """
        Handle an error in the worker thread.
        
        Args:
            error_message (str): Error message
        """
        self.log_message(f"Error: {error_message}")
        
        # Update UI
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.start_btn.setText("Processing... 0%")  # Progress indicator on Start button
        
        # Show error message box
        QMessageBox.critical(
            self,
            "Error",
            f"An error occurred during processing:\n{error_message}"
        )
    
    def closeEvent(self, event):
        """
        Handle window close event.
        
        Args:
            event: Close event
        """
        # Save settings
        self.save_settings()
        
        # Stop worker thread if running
        if self.worker_thread and self.worker_thread.isRunning():
            self.worker_thread.terminate()
            self.worker_thread.wait()
        
        # Accept the close event
        event.accept()

    def open_folder(self, folder_path):
        """
        Open the folder in the file explorer.
        
        Args:
            folder_path (Path): Path to the folder to open
        """
        folder_path = Path(folder_path)
        if not folder_path.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")
        
        # Platform-specific folder opening
        system = platform.system()
        
        try:
            if system == 'Windows':
                os.startfile(str(folder_path))
            elif system == 'Darwin':  # macOS
                subprocess.run(['open', str(folder_path)])
            else:  # Linux and others
                subprocess.run(['xdg-open', str(folder_path)])
        except Exception as e:
            self.logger.error(f"Error opening folder: {e}")
            raise

    def update_progress_percentage(self, percentage):
        """
        Update progress bar with the percentage value.
        
        Args:
            percentage (int): Current percentage (0-100)
        """
        # Progress percentage on Start button
        if percentage == 0:
            self.start_btn.setText("Start Processing")
            self.start_btn.setStyleSheet("""
                QPushButton {
                    background-color: #107C10;
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }
                QPushButton:hover {
                    background-color: #0A6A0A;
                }
                QPushButton:pressed {
                    background-color: #095909;
                }
            """)
        elif percentage == 100:
            # Special style when task is complete
            self.start_btn.setText("✓ Completed")
            self.start_btn.setStyleSheet("""
                QPushButton {
                    background-color: #107C10;
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }
            """)
        else:
            # Button text and background color change
            # Linear gradient for progress indicator
            self.start_btn.setText(f"Processing... {percentage}%")
            
            # Linear gradient for progress indicator
            # Green color increases, darker background color decreases
            self.start_btn.setStyleSheet(f"""
                QPushButton {{
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                                               stop:0 #107C10, 
                                               stop:{percentage/100} #107C10, 
                                               stop:{percentage/100+0.001} #07520A, 
                                               stop:1 #07520A);
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }}
            """)
            
        # Update UI elements
        QApplication.processEvents()

    def save_results(self):
        """
        Save the contents of the results list to a file.
        """
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Results",
            f"litorganizer_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "Text Files (*.txt);;All Files (*)"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("=== LitOrganizer Results ===\n\n")
                    for i in range(self.results_list.count()):
                        item = self.results_list.item(i)
                        f.write(f"{item.text()}\n")
                
                self.logger.info(f"Results saved to: {file_path}")
                QMessageBox.information(self, "Save Successful", f"Results saved to:\n{file_path}")
            except Exception as e:
                self.logger.error(f"Error saving results: {e}")
                QMessageBox.critical(self, "Save Error", f"Failed to save results: {e}")
    
    def save_log(self):
        """
        Save the contents of the log text edit to a file.
        """
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Log",
            f"litorganizer_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "Text Files (*.txt);;All Files (*)"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.log_text.toPlainText())
                
                self.logger.info(f"Log saved to: {file_path}")
                QMessageBox.information(self, "Save Successful", f"Log saved to:\n{file_path}")
            except Exception as e:
                self.logger.error(f"Error saving log: {e}")
                QMessageBox.critical(self, "Save Error", f"Failed to save log: {e}")

    def add_shadow_effect(self, widget, blur_radius=8, color=QColor(0, 0, 0, 25), x_offset=0, y_offset=2):
        """
        Add shadow effect to a widget for modern appearance
        
        Args:
            widget: The widget to add shadow to
            blur_radius: Radius of the shadow blur effect
            color: Color of the shadow
            x_offset: Horizontal offset of the shadow
            y_offset: Vertical offset of the shadow
        """
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(blur_radius)
        shadow.setColor(color)
        shadow.setOffset(x_offset, y_offset)
        widget.setGraphicsEffect(shadow)

    def browse_search_directory(self):
        """
        Browse for a directory containing PDF files to search within.
        """
        directory = QFileDialog.getExistingDirectory(
            self, "Select PDF Directory", str(Path.home())
        )
        if directory:
            self.search_dir_edit.setText(directory)
    
    def start_search_processing(self):
        """
        Start the search processing in a separate thread.
        """
        directory = self.search_dir_edit.text()
        if not directory:
            QMessageBox.warning(self, "Error", "Please select a PDF directory first.")
            return
        
        # Check if directory exists
        if not os.path.exists(directory):
            QMessageBox.warning(self, "Error", "The selected directory doesn't exist.")
            return
        
        # Get keyword
        keyword = self.search_keyword_edit.text().strip()
        if not keyword:
            QMessageBox.warning(self, "Error", "Please enter a keyword to search for.")
            return
        
        # Get search options
        exact_match = self.exact_match_check.isChecked()
        case_sensitive = self.case_match_check.isChecked()
        use_regex = self.regex_match_check.isChecked()
        
        # If using regex, check if it's valid
        if use_regex:
            try:
                re.compile(keyword)
            except re.error:
                QMessageBox.warning(self, "Error", "Invalid regular expression. Please correct it and try again.")
                return
        
        # Clear previous results
        self.clear_search_results()
        
        # Create a worker thread for processing
        self.search_worker = SearchKeywordWorkerThread(
            directory=Path(directory),
            keyword=keyword,
            exact_match=exact_match,
            case_sensitive=case_sensitive,
            use_regex=use_regex,
            logger=self.logger
        )
        
        # Connect signals
        self.search_worker.progress_update.connect(self.log_message)
        self.search_worker.file_processed.connect(self.update_search_file_status)
        self.search_worker.processing_complete.connect(self.search_processing_completed)
        self.search_worker.progress_percentage.connect(self.search_update_progress_percentage)
        self.search_worker.result_found.connect(self.add_search_result)
        self.search_worker.error_occurred.connect(self.handle_error)
        
        # Update UI
        self.search_start_btn.setEnabled(False)
        self.search_stop_btn.setEnabled(True)
        self.search_start_btn.setText("Searching... 0%")  # Progress indicator on Start button
        self.save_search_results_btn.setEnabled(False)
        
        # Start processing
        match_type = []
        if exact_match:
            match_type.append("exact match")
        if case_sensitive:
            match_type.append("case sensitive")
        if use_regex:
            match_type.append("regex")
        match_desc = f" ({', '.join(match_type)})" if match_type else ""
        
        self.log_message(f"🔍 Starting search for: '{keyword}'{match_desc}")
        self.search_worker.start()
        
    def stop_search_processing(self):
        """
        Stop the search processing thread.
        """
        if hasattr(self, 'search_worker') and self.search_worker.isRunning():
            self.search_worker.terminate_flag = True
            self.search_worker.wait()
            self.log_message("⛔ Search stopped by user.")
            
            # Progress bar pause
            self.search_start_btn.setText("Stopping...")
            self.search_start_btn.setStyleSheet("""
                QPushButton {
                    background-color: #F8A800;
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }
            """)
            
        # Update UI
        self.search_start_btn.setEnabled(True)
        self.search_stop_btn.setEnabled(False)
        
        # Reset button text
        self.search_start_btn.setText("Start Search")
        self.search_start_btn.setStyleSheet("""
            QPushButton {
                background-color: #107C10;
                color: white;
                border: none;
                padding: 14px 30px;
                border-radius: 10px;
                font-weight: 600;
                font-size: 14px;
                text-align: center;
            }
            QPushButton:hover {
                background-color: #0A6A0A;
            }
            QPushButton:pressed {
                background-color: #095909;
            }
        """)
    
    def search_update_progress_percentage(self, percentage):
        """
        Update the progress during search processing.
        """
        # Progress percentage on Start button
        if percentage == 0:
            self.search_start_btn.setText("Start Search")
            self.search_start_btn.setStyleSheet("""
                QPushButton {
                    background-color: #107C10;
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }
                QPushButton:hover {
                    background-color: #0A6A0A;
                }
                QPushButton:pressed {
                    background-color: #095909;
                }
            """)
        elif percentage == 100:
            # Special style when task is complete
            self.search_start_btn.setText("✓ Completed")
            self.search_start_btn.setStyleSheet("""
                QPushButton {
                    background-color: #107C10;
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }
            """)
        else:
            # Show progress - Add progress effect
            self.search_start_btn.setText(f"Searching... {percentage}%")
            
            # Gradient for progress effect
            self.search_start_btn.setStyleSheet(f"""
                QPushButton {{
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                        stop:0 #107C10, 
                        stop:{percentage/100} #107C10, 
                        stop:{percentage/100 + 0.001} #084E08, 
                        stop:1 #084E08);
                    color: white;
                    border: none;
                    padding: 14px 30px;
                    border-radius: 10px;
                    font-weight: 600;
                    font-size: 14px;
                    text-align: center;
                }}
            """)
    
    def update_search_file_status(self, filename, success):
        """
        Update the status of a processed file in the search processing.
        """
        if success:
            self.log_message(f"✅ Processed: {filename}")
        else:
            self.log_message(f"❌ Processing failed: {filename}")
    
    def add_search_result(self, doi, filename, page_number, keyword, prev_sentence, matched_sentence, next_sentence, citation=None):
        """
        Add a search result to the results table.
        """
        # Normalize DOI format - remove any prefixes and keep only the actual DOI part (10.xxxx/...)
        if doi:
            # Find the 10. part which indicates the beginning of actual DOI
            doi_start = doi.find("10.")
            if doi_start >= 0:
                doi = doi[doi_start:]
        
        row_position = self.search_results_table.rowCount()
        self.search_results_table.insertRow(row_position)
        
        self.search_results_table.setItem(row_position, 0, QTableWidgetItem(doi or ""))
        self.search_results_table.setItem(row_position, 1, QTableWidgetItem(filename))
        self.search_results_table.setItem(row_position, 2, QTableWidgetItem(str(page_number)))
        self.search_results_table.setItem(row_position, 3, QTableWidgetItem(keyword))
        self.search_results_table.setItem(row_position, 4, QTableWidgetItem(prev_sentence))
        self.search_results_table.setItem(row_position, 5, QTableWidgetItem(matched_sentence))
        self.search_results_table.setItem(row_position, 6, QTableWidgetItem(next_sentence))
    
    def search_processing_completed(self, processed, found):
        """
        Handle completion of search processing.
        """
        # Update UI
        self.search_start_btn.setEnabled(True)
        self.search_stop_btn.setEnabled(False)
        self.search_update_progress_percentage(100)
        
        if found > 0:
            self.save_search_results_btn.setEnabled(True)
            self.log_message(f"✅ Search completed. Found {found} matches in {processed} files.")
        else:
            self.log_message(f"🔍 Search completed. No matches found in {processed} files.")
    
    def clear_search_results(self):
        """
        Clear the search results.
        """
        self.search_results_table.setRowCount(0)
        self.save_search_results_btn.setEnabled(False)
        # Reset progress button
        self.search_start_btn.setText("Start Search")
        self.search_start_btn.setStyleSheet("""
            QPushButton {
                background-color: #107C10;
                color: white;
                border: none;
                padding: 14px 30px;
                border-radius: 10px;
                font-weight: 600;
                font-size: 14px;
                text-align: center;
            }
            QPushButton:hover {
                background-color: #0A6A0A;
            }
            QPushButton:pressed {
                background-color: #095909;
            }
        """)
        self.log_message("🧹 Search results cleared.")
    
    def save_search_results(self):
        """
        Save the search results to Excel and Word files.
        """
        if self.search_results_table.rowCount() == 0:
            QMessageBox.warning(self, "Error", "No search results to save.")
            return
            
        # First, ask user for a filename
        keyword = self.search_keyword_edit.text().strip()
        default_filename = keyword.replace(' ', '_') if keyword else "search_results"
        
        filename_dialog = QDialog(self)
        filename_dialog.setWindowTitle("Enter File Name")
        filename_dialog.setFixedWidth(400)
        filename_dialog.setStyleSheet("""
            QDialog {
                background-color: #FAFAFA;
            }
            QLabel {
                font-size: 10pt;
            }
            QLineEdit {
                padding: 8px;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                font-size: 10pt;
            }
            QPushButton {
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                color: white;
            }
            #save_btn {
                background-color: #107C10;
            }
            #save_btn:hover {
                background-color: #0A6A0A;
            }
            #cancel_btn {
                background-color: #888;
            }
            #cancel_btn:hover {
                background-color: #666;
            }
        """)
        
        layout = QVBoxLayout(filename_dialog)
        
        # Information label
        info_label = QLabel("Enter a name for the result files.\nBoth Excel and Word files will be created with this name.")
        layout.addWidget(info_label)
        
        # Filename input
        filename_input = QLineEdit(default_filename)
        layout.addWidget(filename_input)
        
        # Buttons
        btn_layout = QHBoxLayout()
        save_btn = QPushButton("Save")
        save_btn.setObjectName("save_btn")
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setObjectName("cancel_btn")
        
        btn_layout.addWidget(cancel_btn)
        btn_layout.addWidget(save_btn)
        
        layout.addLayout(btn_layout)
        
        # Button connections
        save_btn.clicked.connect(filename_dialog.accept)
        cancel_btn.clicked.connect(filename_dialog.reject)
        
        if filename_dialog.exec() != QDialog.Accepted:
            return
        
        # User's input filename
        custom_filename = filename_input.text().strip()
        if not custom_filename:
            custom_filename = default_filename
        
        # Clean invalid characters
        import re
        custom_filename = re.sub(r'[<>:"/\\|?*]', '_', custom_filename)
        
        # Ask user for directory to save the results
        directory = QFileDialog.getExistingDirectory(
            self, "Select Folder to Save Results", str(Path.home())
        )
        if not directory:
            return
            
        results_folder = Path(directory)
        
        try:
            # Ensure results directory exists
            if not results_folder.exists():
                results_folder.mkdir(parents=True)
                
            # Create DOCX file with results
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.enum.style import WD_STYLE_TYPE
            
            word_doc = Document()
            word_doc.add_heading(f'Search Results for "{keyword}"', 0)
            
            # Add results to Word document with enhanced formatting
            current_doi = None
            for row in range(self.search_results_table.rowCount()):
                doi = self.search_results_table.item(row, 0).text()
                filename = self.search_results_table.item(row, 1).text()
                page_number = self.search_results_table.item(row, 2).text()
                keyword = self.search_results_table.item(row, 3).text()
                prev_sentence = self.search_results_table.item(row, 4).text()
                matched_sentence = self.search_results_table.item(row, 5).text()
                next_sentence = self.search_results_table.item(row, 6).text()
                
                # Add DOI as heading if it's a new DOI
                if doi and doi != current_doi:
                    # Add DOI as a normal paragraph with bold formatting
                    doi_para = word_doc.add_paragraph()
                    # Ensure the DOI is in the normalized format (starts with 10.)
                    if not doi.startswith("10."):
                        doi_start = doi.find("10.")
                        if doi_start >= 0:
                            doi = doi[doi_start:]
                    
                    doi_para.add_run(f"DOI: {doi}").bold = True
                    current_doi = doi
                
                # Add file information with page number as Heading 1
                file_heading = word_doc.add_heading(level=1)
                file_heading.add_run(f"File: {filename}")
                
                # Add page number as a normal paragraph
                page_para = word_doc.add_paragraph()
                page_para.add_run("Page: ").bold = True
                page_para.add_run(page_number)
                
                # Add matched sentence with context
                context_para = word_doc.add_paragraph()
                
                # Previous sentence (if any)
                if prev_sentence:
                    context_para.add_run(prev_sentence + " ")
                
                # For the matched sentence, we want to highlight the keyword
                # First find where the keyword is in the matched sentence
                if self.regex_match_check.isChecked():
                    # With regex, we can't reliably identify the exact match
                    # So just highlight the entire sentence
                    match_run = context_para.add_run(matched_sentence)
                    match_run.bold = True  
                    match_run.font.highlight_color = 7  # Yellow highlight
                else:
                    # For non-regex searches, try to highlight the exact match
                    import re
                    pattern = re.escape(keyword)
                    flags = 0 if self.case_match_check.isChecked() else re.IGNORECASE
                    
                    # Find all occurrences of the keyword in the matched sentence
                    matches = list(re.finditer(pattern, matched_sentence, flags))
                    
                    if matches:
                        # Add parts of the sentence, highlighting each match
                        last_end = 0
                        for match in matches:
                            # Add text before the match
                            if match.start() > last_end:
                                before_text = matched_sentence[last_end:match.start()]
                                before_run = context_para.add_run(before_text)
                                before_run.font.highlight_color = 7  # Yellow highlight
                            
                            # Add the matched keyword with extra formatting
                            keyword_text = matched_sentence[match.start():match.end()]
                            keyword_run = context_para.add_run(keyword_text)
                            keyword_run.bold = True
                            keyword_run.underline = True
                            keyword_run.font.highlight_color = 7  # Yellow highlight
                            
                            last_end = match.end()
                        
                        # Add any remaining text after the last match
                        if last_end < len(matched_sentence):
                            after_text = matched_sentence[last_end:]
                            after_run = context_para.add_run(after_text)
                            after_run.font.highlight_color = 7  # Yellow highlight
                    else:
                        # Fallback if we can't find the keyword (shouldn't happen)
                        match_run = context_para.add_run(matched_sentence)
                        match_run.bold = True
                        match_run.font.highlight_color = 7  # Yellow highlight
                
                # Next sentence (if any)
                if next_sentence:
                    context_para.add_run(" " + next_sentence)
                
                # Add separator
                word_doc.add_paragraph("-" * 60)
            
            # Save Word document with custom filename
            word_path = results_folder / f"{custom_filename}.docx"
            word_doc.save(word_path)
            
            # Create DataFrame for Excel
            import pandas as pd
            
            data = []
            for row in range(self.search_results_table.rowCount()):
                row_data = []
                for col in range(self.search_results_table.columnCount()):
                    row_data.append(self.search_results_table.item(row, col).text())
                data.append(row_data)
            
            df = pd.DataFrame(
                data, 
                columns=["DOI", "PDF Name", "Page", "Keyword", "Previous Sentence", "Matched Sentence", "Next Sentence"]
            )
            
            # Save to Excel with custom filename
            excel_path = results_folder / f"{custom_filename}.xlsx"
            df.to_excel(excel_path, index=False)
            
            # Show success message
            QMessageBox.information(
                self,
                "Success",
                f"Results saved successfully:\n\n"
                f"Excel: {excel_path}\n"
                f"Word: {word_path}"
            )
            
            # Log the message
            self.log_message(f"✅ Results saved to: {results_folder}")
            
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to save results: {str(e)}"
            )
            self.logger.error(f"Failed to save search results: {e}")
            self.logger.debug(traceback.format_exc())

    def show_search_options_help(self):
        """
        Show detailed help for search options in a modal dialog.
        """
        help_dialog = QDialog(self)
        help_dialog.setWindowTitle("Search Options Help")
        help_dialog.setMinimumWidth(600)
        help_dialog.setMinimumHeight(400)
        help_dialog.setWindowIcon(QIcon("resources/question.png"))
        help_dialog.setStyleSheet("""
            QDialog {
                background-color: #FAFAFA;
            }
            QLabel {
                font-size: 10pt;
                margin-bottom: 10px;
            }
            QLabel#title {
                font-size: 14pt;
                font-weight: bold;
                color: #0078D4;
                margin-bottom: 15px;
            }
            QLabel#subtitle {
                font-size: 12pt;
                font-weight: bold;
                color: #333;
                margin-top: 10px;
                margin-bottom: 5px;
            }
            QPushButton {
                padding: 8px 16px;
                background-color: #0078D4;
                color: white;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #0065B8;
            }
        """)
        
        layout = QVBoxLayout(help_dialog)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(10)
        
        # Title
        title = QLabel("Search Options Explained")
        title.setObjectName("title")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)
        
        # Scrollable content area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: transparent;
            }
        """)
        
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(10, 10, 10, 10)
        scroll_layout.setSpacing(15)
        
        # Introduction
        intro = QLabel(
            "The Search Keywords feature allows you to find specific text within your PDF documents. "
            "You can customize your search using the following options:"
        )
        intro.setWordWrap(True)
        scroll_layout.addWidget(intro)
        
        # Multi-word terms explanation
        multi_subtitle = QLabel("Searching for Multiple Words or Phrases")
        multi_subtitle.setObjectName("subtitle")
        scroll_layout.addWidget(multi_subtitle)
        
        multi_explanation = QLabel(
            "<b>Basic Behavior:</b> By default, when you search for multiple words (e.g., 'deformation analysis'), "
            "the search will find any instances where these words appear together in that exact order.<br><br>"
            "<b>Examples:</b><br>"
            "• Searching for 'deformation analysis' will find the exact phrase 'deformation analysis'<br>"
            "• It will NOT find instances where the words are separate (e.g., 'deformation in the analysis')<br>"
            "• It will NOT find variations like 'deformational analysis' unless exact match is turned off<br><br>"
            "<b>Strategies for Multi-word Terms:</b><br>"
            "• <b>Without Exact Match:</b> Searching for 'deformation analysis' will find partial matches like 'deformational analysis'<br>"
            "• <b>With Regex:</b> Use '\\bdeformation\\s+analysis\\b' to find the words with any whitespace between them<br>"
            "• <b>Flexible Order:</b> Use regex with '(?=.*deformation)(?=.*analysis)' to find pages containing both words in any order<br><br>"
            "<b>Tip:</b> For complex phrases, turning off 'Exact Match' may yield better results by allowing for word variations."
        )
        multi_explanation.setWordWrap(True)
        scroll_layout.addWidget(multi_explanation)
        
        # Exact Match explanation
        exact_subtitle = QLabel("Exact Match")
        exact_subtitle.setObjectName("subtitle")
        scroll_layout.addWidget(exact_subtitle)
        
        exact_explanation = QLabel(
            "<b>What it does:</b> When enabled, the search will only find complete words that exactly match your keyword.<br><br>"
            "<b>Examples:</b><br>"
            "• Searching for 'test' <u>will</u> find: 'This is a test.' or 'The test results are in.'<br>"
            "• Searching for 'test' <u>will not</u> find: 'testing', 'attest', or 'contest'<br><br>"
            "<b>When to use:</b> Use this option when you want to find only the exact word and avoid partial matches. "
            "This is particularly useful for common word fragments that might appear in many different words."
        )
        exact_explanation.setWordWrap(True)
        scroll_layout.addWidget(exact_explanation)
        
        # Case Sensitive explanation
        case_subtitle = QLabel("Case Sensitive")
        case_subtitle.setObjectName("subtitle")
        scroll_layout.addWidget(case_subtitle)
        
        case_explanation = QLabel(
            "<b>What it does:</b> When enabled, the search will match the exact letter case (uppercase/lowercase) of your keyword.<br><br>"
            "<b>Examples:</b><br>"
            "• Searching for 'Test' with case sensitivity <u>will</u> find: 'Test case' or 'The Test showed...'<br>"
            "• Searching for 'Test' with case sensitivity <u>will not</u> find: 'test', 'TEST', or 'tEsT'<br><br>"
            "<b>When to use:</b> Use this option when the specific capitalization matters. "
            "This is useful for proper nouns, acronyms, or technical terms where capitalization is significant."
        )
        case_explanation.setWordWrap(True)
        scroll_layout.addWidget(case_explanation)
        
        # Regex explanation
        regex_subtitle = QLabel("Use Regex (Regular Expressions)")
        regex_subtitle.setObjectName("subtitle")
        scroll_layout.addWidget(regex_subtitle)
        
        regex_explanation = QLabel(
            "<b>What it does:</b> Enables the use of regular expressions for powerful pattern matching.<br><br>"
            "<b>Example patterns:</b><br>"
            "• <code>word|term</code> - Finds either 'word' or 'term'<br>"
            "• <code>test.*data</code> - Finds 'test' followed by any characters and then 'data'<br>"
            "• <code>\\d+</code> - Finds one or more digits<br>"
            "• <code>[A-Z]\\w+</code> - Finds words that start with a capital letter<br><br>"
            "<b>When to use:</b> Use regular expressions for advanced search patterns that can't be expressed with simple text. "
            "This is powerful for finding complex patterns, but requires knowledge of regex syntax.<br><br>"
            "<b>Note:</b> When 'Use Regex' is enabled, the 'Exact Match' option is only applied if your pattern doesn't already "
            "include word boundary markers (<code>\\b</code>)."
        )
        regex_explanation.setWordWrap(True)
        scroll_layout.addWidget(regex_explanation)
        
        # Combining options
        combining_subtitle = QLabel("Combining Search Options")
        combining_subtitle.setObjectName("subtitle")
        scroll_layout.addWidget(combining_subtitle)
        
        combining_explanation = QLabel(
            "You can combine these options for more precise searches:<br><br>"
            "• <b>Exact Match + Case Sensitive:</b> Finds only complete words with the exact same capitalization<br>"
            "• <b>Regex + Case Sensitive:</b> Applies your regex pattern with case sensitivity<br><br>"
            "Experiment with different combinations to find the most effective search approach for your documents."
        )
        combining_explanation.setWordWrap(True)
        scroll_layout.addWidget(combining_explanation)
        
        scroll_area.setWidget(scroll_content)
        layout.addWidget(scroll_area)
        
        # Close button at bottom
        button_layout = QHBoxLayout()
        close_button = QPushButton("Close")
        close_button.clicked.connect(help_dialog.accept)
        button_layout.addStretch(1)
        button_layout.addWidget(close_button)
        button_layout.addStretch(1)
        layout.addLayout(button_layout)
        
        # Show the dialog
        help_dialog.exec_()

    def update_statistics(self, processed, renamed, problematic):
        """
        Update statistics tab with the processing results.
        
        Args:
            processed (int): Number of processed files
            renamed (int): Number of renamed files
            problematic (int): Number of unclassified files
        """
        # Calculate success rate
        success_rate = round((renamed / processed) * 100 if processed > 0 else 0, 1)
        
        # Update general statistics
        self.stat_total_files.setText(f"Total Files Processed: {processed}")
        self.stat_renamed_files.setText(f"Successfully Renamed: {renamed}")
        self.stat_problematic_files.setText(f"Files Not Renamed: {problematic}")
        self.stat_success_rate.setText(f"Success Rate: {success_rate}%")
        
        # Calculate and update performance metrics
        if self.process_start_time and self.process_end_time:
            total_time = round(self.process_end_time - self.process_start_time, 2)
            time_per_file = round(total_time / processed, 2) if processed > 0 else 0
            processing_speed = round(processed / total_time, 2) if total_time > 0 else 0
            
            self.stat_total_time.setText(f"Total Processing Time: {total_time} seconds")
            self.stat_per_file_time.setText(f"Average Time Per File: {time_per_file} seconds")
            self.stat_processing_speed.setText(f"Processing Speed: {processing_speed} files/second")
            
            # Estimate memory usage (simplified)
            estimated_memory = round(processed * 0.5, 1)  # Rough estimate, 0.5 MB per file
            self.stat_memory_usage.setText(f"Estimated Memory Usage: {estimated_memory} MB")
            
            # Estimate time savings (simplified)
            # Assume manual renaming takes 30 seconds per file
            manual_time = processed * 30
            time_saved = manual_time - total_time
            time_saved_minutes = round(time_saved / 60, 1)
            self.stat_time_savings.setText(f"Estimated Time Saved: {time_saved_minutes} minutes")
            
            # Update progress bars for performance metrics
            # Scale them to reasonable values
            max_expected_time = 300  # 5 minutes
            self.total_time_progress.setValue(min(int(total_time / max_expected_time * 100), 100))
            
            max_expected_per_file_time = 10  # 10 seconds
            self.per_file_time_progress.setValue(min(int(time_per_file / max_expected_per_file_time * 100), 100))
            
            max_expected_speed = 5  # 5 files/second
            self.processing_speed_progress.setValue(min(int(processing_speed / max_expected_speed * 100), 100))
            
            max_expected_memory = 100  # 100 MB
            self.memory_usage_progress.setValue(min(int(estimated_memory / max_expected_memory * 100), 100))
            
            # Update time savings progress bar
            max_expected_time_saved = 60  # 60 minutes (1 hour)
            self.time_savings_progress.setValue(min(int(time_saved_minutes / max_expected_time_saved * 100), 100))
        
        # Update accuracy metrics based on success rate and log analysis
        log_text = self.log_text.toPlainText()
        
        # Calculate DOI detection rate (based on how many files were successfully processed)
        doi_detection_rate = success_rate  # As a simplification, use success rate
        
        # Calculate metadata quality based on API source distribution 
        metadata_quality = round(success_rate * 0.9, 1)  # Simplified estimate
        
        # Calculate categorization quality based on how many categorization options were enabled
        categorization_options_enabled = sum([
            self.categorize_by_journal_check.isChecked(),
            self.categorize_by_author_check.isChecked(),
            self.categorize_by_year_check.isChecked(),
            self.categorize_by_subject_check.isChecked()
        ])
        
        categorization_quality = round(100 * (categorization_options_enabled / 4) if categorization_options_enabled > 0 else 0, 1)
        
        # Update the accuracy metrics
        self.stat_doi_detection.setText(f"DOI Detection Rate: {doi_detection_rate}%")
        self.stat_metadata_quality.setText(f"Metadata Quality: {metadata_quality}%")
        self.stat_categorization_quality.setText(f"Categorization Quality: {categorization_quality}%")
        
        # Update accuracy metrics progress bars
        self.doi_detection_progress.setValue(int(doi_detection_rate))
        self.metadata_quality_progress.setValue(int(metadata_quality))
        self.categorization_quality_progress.setValue(int(categorization_quality))
        
        # Hide the placeholder message
        self.stats_placeholder.setVisible(False)
        
        # Clear previous API stats
        self.api_stats_list.clear()
        
        # Collect metadata sources from log messages
        api_sources = {}
        
        # Search for API source mentions in the log
        api_matches = re.findall(r"Found metadata using ([a-zA-Z0-9]+) for", log_text)
        for api in api_matches:
            api_sources[api] = api_sources.get(api, 0) + 1
        
        # Add API stats to the list
        if api_sources:
            for api, count in sorted(api_sources.items(), key=lambda x: x[1], reverse=True):
                percentage = round((count / renamed) * 100 if renamed > 0 else 0, 1)
                item = QListWidgetItem(f"{api}: {count} files ({percentage}%)")
                self.api_stats_list.addItem(item)
        else:
            self.api_stats_list.addItem("No API source information found")
        
        # Update categorization statistics if options are enabled
        if self.categorize_by_journal_check.isChecked():
            self.stat_categorized_by_journal.setText(f"Categorized by Journal: {renamed} files")
        else:
            self.stat_categorized_by_journal.setText("Categorized by Journal: Disabled")
        
        if self.categorize_by_author_check.isChecked():
            self.stat_categorized_by_author.setText(f"Categorized by Author: {renamed} files")
        else:
            self.stat_categorized_by_author.setText("Categorized by Author: Disabled")
        
        if self.categorize_by_year_check.isChecked():
            self.stat_categorized_by_year.setText(f"Categorized by Year: {renamed} files")
        else:
            self.stat_categorized_by_year.setText("Categorized by Year: Disabled")
        
        if self.categorize_by_subject_check.isChecked():
            self.stat_categorized_by_subject.setText(f"Categorized by Subject: {renamed} files")
        else:
            self.stat_categorized_by_subject.setText("Categorized by Subject: Disabled")
        
        # Update error breakdown
        self.error_breakdown_list.clear()
        log_text = self.log_text.toPlainText() # Get the full log text

        if problematic > 0:
            # --- Define regex patterns for specific error types ---
            # Note: These patterns should match the specific log messages added in pdf_renamer.py
            # We use r"..." for raw strings to handle backslashes in patterns if needed.
            # Using re.IGNORECASE for flexibility, although error logs are likely consistent.

            # Core known issues
            no_doi_pattern = r"No DOI found in"
            no_metadata_pattern = r"Insufficient or no metadata found for DOI" # Updated pattern

            # New specific categories based on refined logging
            pdf_read_pattern = r"PDF Processing Error \(read\)" # Escaped parentheses
            pdf_encrypt_pattern = r"PDF Processing Error \(encrypted\)" # Escaped parentheses
            doi_extract_pattern = r"Error extracting DOI from"
            api_error_pattern = r"API Error \(network/http\)" # Escaped parentheses
            metadata_fetch_pattern = r"Error fetching metadata for DOI" # General metadata fetch error
            fs_error_pattern = r"File System Error \((mkdir|copy|unlink|backup|rename/move|categorize move)\)" # Escaped parentheses
            rename_move_error_pattern = r"Error renaming/moving file" # Catch-all for rename/move if FS specific didn't catch
            categorize_error_pattern = r"Error during file categorization attempt|Error calling categorize_file" # Catches categorization errors
            unexpected_error_pattern = r"Unexpected Error processing file"


            # --- Count occurrences ---
            no_doi_count = len(re.findall(no_doi_pattern, log_text, re.IGNORECASE))
            no_metadata_count = len(re.findall(no_metadata_pattern, log_text, re.IGNORECASE))
            pdf_read_count = len(re.findall(pdf_read_pattern, log_text, re.IGNORECASE))
            pdf_encrypt_count = len(re.findall(pdf_encrypt_pattern, log_text, re.IGNORECASE))
            doi_extract_count = len(re.findall(doi_extract_pattern, log_text, re.IGNORECASE))
            api_error_count = len(re.findall(api_error_pattern, log_text, re.IGNORECASE))
            metadata_fetch_count = len(re.findall(metadata_fetch_pattern, log_text, re.IGNORECASE))
            fs_error_count = len(re.findall(fs_error_pattern, log_text, re.IGNORECASE))
            rename_move_count = len(re.findall(rename_move_error_pattern, log_text, re.IGNORECASE))
            categorize_error_count = len(re.findall(categorize_error_pattern, log_text, re.IGNORECASE))
            unexpected_error_count = len(re.findall(unexpected_error_pattern, log_text, re.IGNORECASE))

            # Calculate 'Other' errors - start with total problematic and subtract specifics
            # Be careful not to double-count (e.g., a specific FS error might also contain "Error renaming/moving")
            # Let's sum up all specific counts found
            counted_specific_errors = (
                no_doi_count + no_metadata_count + pdf_read_count + pdf_encrypt_count +
                doi_extract_count + api_error_count + metadata_fetch_count + fs_error_count +
                rename_move_count + categorize_error_count + unexpected_error_count
            )

            # Other errors are those problematic files NOT accounted for by specific patterns
            # It might be slightly inaccurate if logs are complex, but better than before.
            other_errors = max(0, problematic - counted_specific_errors) # Ensure non-negative


            # --- Add items to the list ---
            def add_error_item(label, count):
                if count > 0:
                    percentage = round(count / problematic * 100, 1)
                    self.error_breakdown_list.addItem(f"{label}: {count} files ({percentage}%)")

            add_error_item("Missing DOI", no_doi_count)
            add_error_item("Insufficient/No Metadata", no_metadata_count) # Updated label
            add_error_item("PDF Read Error", pdf_read_count)
            add_error_item("PDF Encrypted", pdf_encrypt_count)
            add_error_item("DOI Extraction Error", doi_extract_count)
            add_error_item("API Network/HTTP Error", api_error_count)
            add_error_item("Metadata Fetch Error", metadata_fetch_count)
            add_error_item("File System Error", fs_error_count)
            add_error_item("Rename/Move Error", rename_move_count) # May overlap with FS, shows general rename fail
            add_error_item("Categorization Error", categorize_error_count)
            add_error_item("Unexpected Processing Error", unexpected_error_count)

            # Add the remaining 'Other' errors
            if other_errors > 0:
                 add_error_item("Other (Unclassified)", other_errors)

        else:
            # If no problematic files, show success message
            self.error_breakdown_list.addItem("✅ No errors encountered during processing.")

        # --- Update publication statistics ---
        # Year distribution
        self.year_distribution_list.clear()
        
        # Search for year information in the log
        year_stats = {}
        year_matches = re.findall(r"Categorized .+ by year: (\d{4})", log_text)
        
        for year in year_matches:
            year_stats[year] = year_stats.get(year, 0) + 1
        
        # Add year stats to the list
        if year_stats:
            for year, count in sorted(year_stats.items(), key=lambda x: x[0], reverse=True):
                percentage = round((count / renamed) * 100 if renamed > 0 else 0, 1)
                item = QListWidgetItem(f"{year}: {count} files ({percentage}%)")
                self.year_distribution_list.addItem(item)
        elif self.categorize_by_year_check.isChecked():
            self.year_distribution_list.addItem("No files categorized by year")
        else:
            self.year_distribution_list.addItem("Year categorization disabled")
        
        # Author statistics
        self.top_authors_list.clear()
        
        # Extract author information
        author_stats = {}
        author_matches = re.findall(r"Categorized .+ by author: (.+)", log_text)
        
        for author in author_matches:
            author = author.strip()
            author_stats[author] = author_stats.get(author, 0) + 1
        
        # Add all authors to the list (not just top 5)
        if author_stats:
            for author, count in sorted(author_stats.items(), key=lambda x: x[1], reverse=True):
                percentage = round((count / renamed) * 100 if renamed > 0 else 0, 1)
                item = QListWidgetItem(f"{author}: {count} files ({percentage}%)")
                # Set background color for visual ranking (white)
                self.top_authors_list.addItem(item)
        elif self.categorize_by_author_check.isChecked():
            self.top_authors_list.addItem("No author data available")
        else:
            self.top_authors_list.addItem("Author categorization disabled")
        
        # Journal statistics
        self.top_journals_list.clear()
        
        # Extract journal information
        journal_stats = {}
        journal_matches = re.findall(r"Categorized .+ by journal: (.+)", log_text)
        
        for journal in journal_matches:
            journal = journal.strip()
            journal_stats[journal] = journal_stats.get(journal, 0) + 1
        
        # Add all journals to the list (not just top 5)
        if journal_stats:
            for journal, count in sorted(journal_stats.items(), key=lambda x: x[1], reverse=True):
                percentage = round((count / renamed) * 100 if renamed > 0 else 0, 1)
                item = QListWidgetItem(f"{journal}: {count} files ({percentage}%)")
                self.top_journals_list.addItem(item)
        elif self.categorize_by_journal_check.isChecked():
            self.top_journals_list.addItem("No journal data available")
        else:
            self.top_journals_list.addItem("Journal categorization disabled")
        
        # Subject statistics
        self.subject_list.clear()
        
        # Extract subject information
        subject_stats = {}
        subject_matches = re.findall(r"Categorized .+ by subject: (.+)", log_text)
        
        for subject in subject_matches:
            subject = subject.strip()
            subject_stats[subject] = subject_stats.get(subject, 0) + 1
        
        # Add subjects to the list
        if subject_stats:
            for subject, count in sorted(subject_stats.items(), key=lambda x: x[1], reverse=True):
                percentage = round((count / renamed) * 100 if renamed > 0 else 0, 1)
                item = QListWidgetItem(f"{subject}: {count} files ({percentage}%)")
                self.subject_list.addItem(item)
        elif self.categorize_by_subject_check.isChecked():
            self.subject_list.addItem("No subject data available")
        else:
            self.subject_list.addItem("Subject categorization disabled")


class SearchKeywordWorkerThread(QThread):
    """
    Worker thread for searching keywords in PDF files.
    """
    
    # Signals for progress reporting
    progress_update = pyqtSignal(str)
    file_processed = pyqtSignal(str, bool)
    processing_complete = pyqtSignal(int, int)
    error_occurred = pyqtSignal(str)
    progress_percentage = pyqtSignal(int)
    result_found = pyqtSignal(str, str, int, str, str, str, str)  # DOI, filename, page_number, keyword, prev, matched, next
    
    def __init__(
        self,
        directory: Path,
        keyword: str,
        exact_match: bool = False,
        case_sensitive: bool = False,
        use_regex: bool = False,
        logger: Optional[logging.Logger] = None
    ):
        """
        Initialize the search keyword worker thread.
        
        Args:
            directory (Path): Path to the PDF directory
            keyword (str): Keyword to search for in PDFs
            exact_match (bool): Whether to match the exact keyword or partial matches
            case_sensitive (bool): Whether the search should be case sensitive
            use_regex (bool): Whether to use regular expressions for matching
            logger (Optional[logging.Logger]): Logger instance
        """
        super().__init__()
        self.directory = directory
        self.keyword = keyword
        self.exact_match = exact_match
        self.case_sensitive = case_sensitive
        self.use_regex = use_regex
        self.logger = logger or logging.getLogger('litorganizer.searcher')
        self.terminate_flag = False
        
        # File counters
        self.total_files = 0
        self.processed_files = 0
        self.found_matches = 0
    
    def run(self):
        """
        Run the search keyword worker thread.
        """
        try:
            # Log start message
            self.logger.debug(f"SearchKeywordWorkerThread started: Directory={self.directory}, Keyword={self.keyword}, ExactMatch={self.exact_match}, CaseSensitive={self.case_sensitive}, UseRegex={self.use_regex}")
            
            import fitz  # PyMuPDF
            import re
            import traceback
            import gc  # Garbage collector
            from concurrent.futures import ThreadPoolExecutor, as_completed
            from modules.utils.pdf_metadata_extractor import extract_doi
            
            # Memory cleanup
            gc.collect()
            
            # Helper functions
            def extract_text_from_pdf(pdf_path):
                """Extract text from PDF file and add debug info."""
                try:
                    with fitz.open(pdf_path) as doc:
                        self.progress_update.emit(f"✅ File opened: {pdf_path}")
                        # Extract text from PDF
                        result = [page.get_text("text") for page in doc]
                        return result
                except Exception as e:
                    error_msg = f"❌ Error: Could not open {pdf_path}! -> {e}"
                    self.error_occurred.emit(error_msg)
                    self.logger.error(error_msg)
                    self.logger.error(traceback.format_exc())
                    return None

            def clean_text(text):
                """Clean unnecessary characters."""
                try:
                    return re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
                except Exception as e:
                    self.logger.error(f"Text cleaning error: {e}")
                    return text

            def split_sentences(text):
                """Split text into sentences and filter empty strings."""
                try:
                    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|!)\s', text)
                    return [s.strip() for s in sentences if s.strip()]
                except Exception as e:
                    self.logger.error(f"Sentence splitting error: {e}")
                    return [text] if text.strip() else []

            def extract_doi_from_pdf(pdf_path):
                """Extract DOI from PDF using the parser utility."""
                try:
                    doi = extract_doi(pdf_path, False)  # Second parameter is use_ocr
                    if doi:
                        self.progress_update.emit(f"📝 DOI found: {doi}")
                    return doi
                except Exception as e:
                    self.logger.error(f"DOI extraction error: {e}")
                    return None

            def find_keyword_paragraphs(text_list, keyword, filename, doi=None):
                """Find sentences with keywords and their context."""
                results = []
                try:
                    # Prepare pattern based on search options
                    if self.use_regex:
                        # Use keyword as regex pattern directly
                        keyword_pattern = keyword
                    else:
                        # Escape special characters and apply exact match if needed
                        keyword_pattern = r'\b' + re.escape(keyword) + r'\b' if self.exact_match else re.escape(keyword)
                    
                    flags = 0 if self.case_sensitive else re.IGNORECASE
                    
                    for page_num, text in enumerate(text_list, start=1):
                        if self.terminate_flag:
                            return results
                            
                        if not text:
                            continue
                            
                        text = clean_text(text)
                        sentences = split_sentences(text)

                        for i, sentence in enumerate(sentences):
                            if self.terminate_flag:
                                return results
                                
                            try:
                                if re.search(keyword_pattern, sentence, flags):
                                    prev_sentences = " ".join(sentences[max(i-1, 0):i]) if i > 0 else ""
                                    next_sentences = " ".join(sentences[i+1:min(i+2, len(sentences))]) if i+1 < len(sentences) else ""
                                    
                                    # Signal the result found
                                    self.result_found.emit(
                                        doi or "",        # DOI
                                        filename,         # PDF name
                                        page_num,         # Page number
                                        keyword,          # Keyword
                                        prev_sentences,   # Previous sentence
                                        sentence,         # Matched sentence
                                        next_sentences    # Next sentence
                                    )
                                    
                                    # Add to results count
                                    self.found_matches += 1
                            except Exception as e:
                                self.logger.error(f"Matching error: {e} (sentence: {sentence[:50]}...)")
                                continue
                except Exception as e:
                    self.logger.error(f"find_keyword_paragraphs error: {e}")
                    self.logger.error(traceback.format_exc())
                
                return results

            def process_pdf(file_path):
                """Process each PDF file to find matches."""
                if self.terminate_flag:
                    return
                
                try:
                    filename = file_path.name
                    self.progress_update.emit(f"🔍 Processing: {filename}")
                    self.logger.debug(f"Processing PDF: {filename}")
                    
                    # Extract DOI
                    doi = extract_doi_from_pdf(file_path)
                    
                    # Extract text and find matches
                    text_list = extract_text_from_pdf(file_path)
                    if text_list is None:
                        self.file_processed.emit(filename, False)
                        return
                    
                    find_keyword_paragraphs(text_list, self.keyword, filename, doi)
                    self.file_processed.emit(filename, True)
                    
                    # Clean up memory after processing
                    del text_list
                    gc.collect()
                    
                except Exception as e:
                    error_msg = f"❌ PDF processing error: {file_path.name} -> {e}"
                    self.error_occurred.emit(error_msg)
                    self.logger.error(error_msg)
                    self.logger.error(traceback.format_exc())
                    self.file_processed.emit(file_path.name, False)
            
            # Find PDF files
            try:
                pdf_files = list(self.directory.glob("**/*.pdf"))
                self.total_files = len(pdf_files)
                
                if self.total_files == 0:
                    self.progress_update.emit(f"❌ No PDF files found in directory! {self.directory}")
                    self.processing_complete.emit(0, 0)
                    return
                
                self.progress_update.emit(f"📚 Found {self.total_files} PDF files.")
                self.logger.info(f"Will process {self.total_files} PDF files.")
            except Exception as e:
                error_msg = f"❌ Error listing PDF files: {e}"
                self.error_occurred.emit(error_msg)
                self.logger.error(error_msg)
                self.processing_complete.emit(0, 0)
                return
            
            # Process files with ThreadPoolExecutor
            max_workers = min(os.cpu_count() or 4, 4)  # Use at most 4 threads
            
            try:
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    # Submit all PDF files for processing
                    future_to_file = {executor.submit(process_pdf, f): f for f in pdf_files}
                    
                    for future in as_completed(future_to_file):
                        if self.terminate_flag:
                            executor.shutdown(wait=False)
                            break
                        
                        try:
                            # Update progress
                            self.processed_files += 1
                            percentage = int(100 * self.processed_files / self.total_files)
                            self.progress_percentage.emit(percentage)
                            
                            # Clean up memory periodically
                            if self.processed_files % 10 == 0:
                                gc.collect()
                                
                        except Exception as e:
                            error_msg = f"❌ Thread processing error: {e}"
                            self.error_occurred.emit(error_msg)
                            self.logger.error(error_msg)
                            self.logger.error(traceback.format_exc())
                            continue
            except Exception as e:
                error_msg = f"❌ ThreadPoolExecutor error: {e}"
                self.error_occurred.emit(error_msg)
                self.logger.error(error_msg)
                self.logger.error(traceback.format_exc())
            
            # Final memory cleanup
            gc.collect()
            
            # Report processing results
            if self.found_matches > 0:
                self.progress_update.emit(f"✅ Search completed. Found {self.found_matches} matches in {self.processed_files} files.")
                self.logger.info(f"Search completed: {self.found_matches} matches in {self.processed_files} files.")
            else:
                self.progress_update.emit(f"🔍 Search completed. No matches found.")
                self.logger.info("Search completed: No matches found.")
            
            # Send the complete signal
            self.processing_complete.emit(self.processed_files, self.found_matches)
            
        except Exception as e:
            error_msg = f"❌ Critical error in search process: {str(e)}"
            self.error_occurred.emit(error_msg)
            self.logger.error(error_msg)
            self.logger.error(traceback.format_exc())
            
            # Even on error, send the complete signal
            self.processing_complete.emit(self.processed_files, 0)
    
    def terminate(self):
        """
        Terminate the worker thread.
        """
        self.terminate_flag = True
        self.wait()


def launch_gui(logger: Optional[logging.Logger] = None):
    """
    Launch the GUI application.
    
    Args:
        logger (Optional[logging.Logger]): Logger instance
    """
    # Add file handler for detailed logging
    try:
        import datetime
        import sys
        import os
        from pathlib import Path
        
        # Create log directory
        log_dir = Path("logs")
        if not log_dir.exists():
            log_dir.mkdir(exist_ok=True)
            
        # Log file name with date and time
        current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = log_dir / f"search_term_{current_time}.log"
        
        # Add file handler to root logger
        if logger is None:
            logger = logging.getLogger('litorganizer')
            
        # Set log level to DEBUG
        logger.setLevel(logging.DEBUG)
        
        # File handler for logging
        file_handler = logging.FileHandler(log_file, 'w', 'utf-8')
        file_handler.setLevel(logging.DEBUG)
        
        # Log format
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        
        # Add handler to logger
        logger.addHandler(file_handler)
        
        # Handler for uncaught exceptions
        def exception_hook(exctype, value, traceback_obj):
            """
            Log uncaught exceptions and notify user.
            """
            import traceback
            logger.critical("Unexpected error:", exc_info=(exctype, value, traceback_obj))
            
            # Write error message to file
            error_file = log_dir / f"error_{current_time}.txt"
            with open(error_file, 'w', encoding='utf-8') as f:
                f.write(f"Date and time: {datetime.datetime.now()}\n\n")
                f.write("Error:\n")
                traceback.print_exception(exctype, value, traceback_obj, file=f)
                
            # Call original exception handler
            sys.__excepthook__(exctype, value, traceback_obj)
            
        # Set global exception handler
        sys.excepthook = exception_hook
        
        logger.info("Starting in GUI mode with enhanced logging")
        logger.info(f"Log file: {log_file}")
        
    except Exception as e:
        print(f"Error creating log file: {e}")
        if logger is None:
            logger = logging.getLogger('litorganizer')
    
    # Create QApplication instance
    app = QApplication(sys.argv)
    app.setStyle("Fusion")  # Modern look across platforms
    
    # Catch application crashes
    try:
        # Create and show the main window
        window = MainWindow(logger)
        window.show()
        
        # Start the event loop
        sys.exit(app.exec_())
    except Exception as e:
        logger.critical(f"Critical error starting GUI: {e}", exc_info=True)
        from PyQt5.QtWidgets import QMessageBox
        QMessageBox.critical(
            None, 
            "Critical Error",
            f"Application encountered an unexpected error:\n{str(e)}\n\nDetails logged to: {log_file}"
        )
        sys.exit(1)